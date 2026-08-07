"""DocFlow — DOC/DOCX ↔ PDF 双向转换服务

Flask 后端，提供文件上传、Word COM 转换、进度轮询、PDF/DOCX 下载等 API。
支持 DOC/DOCX → PDF 和 PDF → DOCX 两种转换方向。
历史记录使用 SQLite 持久化存储。
"""

import os
import sys
import atexit
import uuid
import time
import json
import re
import shutil
import sqlite3
import subprocess
import threading
import zipfile
import io
import logging
from contextlib import contextmanager
from dataclasses import dataclass
from datetime import datetime

logger = logging.getLogger(__name__)

from flask import Flask, request, jsonify, send_file, send_from_directory

# OCR engine (lazy import to avoid startup cost)
_ocr_engine = None


def _get_ocr_engine():
    global _ocr_engine
    if _ocr_engine is None:
        try:
            from ocr_engine import get_ocr_engine
            _ocr_engine = get_ocr_engine()
        except Exception:
            _ocr_engine = False  # Mark as unavailable
    return _ocr_engine if _ocr_engine is not False else None


def _check_ocr_engines():
    """Check available OCR engines and return status dict."""
    engine = _get_ocr_engine()
    tess_cmd = _find_tesseract()
    return {
        "rapidocr": engine is not None,
        "tesseract": tess_cmd is not None,
        "primary": "RapidOCR" if engine is not None else ("Tesseract" if tess_cmd is not None else None),
        "chinese_optimized": engine is not None,
    }


app = Flask(__name__, static_folder=None)

# ── Directories ──
# When frozen by PyInstaller, bundled read-only assets live in sys._MEIPASS,
# while writable runtime data (uploads, output, db) sits next to the .exe.
if getattr(sys, 'frozen', False):
    BUNDLE_DIR = sys._MEIPASS
    BASE_DIR = os.path.dirname(sys.executable)
else:
    BUNDLE_DIR = os.path.dirname(os.path.abspath(__file__))
    BASE_DIR = BUNDLE_DIR
UPLOAD_DIR = os.path.join(BASE_DIR, "uploads")
OUTPUT_DIR = os.path.join(BASE_DIR, "output")
TEMP_DIR = os.path.join(BASE_DIR, "temp")
FONT_DIR = os.path.join(BASE_DIR, "fonts")
MUSIC_DIR = os.path.join(BASE_DIR, "music")
DB_PATH = os.path.join(BASE_DIR, "docflow.db")
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(TEMP_DIR, exist_ok=True)
os.makedirs(FONT_DIR, exist_ok=True)
os.makedirs(MUSIC_DIR, exist_ok=True)
_output_lock = threading.Lock()


def _cleanup_dir(target_dir, remove_dirs=False):
    """Remove all files (and optionally subdirectories) in target_dir."""
    if not os.path.isdir(target_dir):
        return
    for name in os.listdir(target_dir):
        path = os.path.join(target_dir, name)
        try:
            if os.path.isfile(path):
                os.remove(path)
            elif os.path.isdir(path) and remove_dirs:
                shutil.rmtree(path)
        except OSError:
            pass


atexit.register(lambda: _cleanup_dir(TEMP_DIR, remove_dirs=True))
atexit.register(lambda: _cleanup_dir(UPLOAD_DIR))


# ── DB context manager ──

@contextmanager
def get_db():
    """Context manager that yields a SQLite connection and closes it on exit."""
    db = sqlite3.connect(DB_PATH, check_same_thread=False)
    db.row_factory = sqlite3.Row
    db.execute("PRAGMA journal_mode=WAL")
    try:
        yield db
    finally:
        db.close()


# ── SQLite helpers ──


def init_db():
    """Create tables and reset stale jobs."""
    with get_db() as db:
        db.execute("""
            CREATE TABLE IF NOT EXISTS jobs (
                id TEXT PRIMARY KEY,
                name TEXT NOT NULL,
                status TEXT NOT NULL DEFAULT 'pending',
                progress INTEGER DEFAULT 0,
                input_path TEXT,
                output_path TEXT,
                input_size INTEGER DEFAULT 0,
                input_size_formatted TEXT DEFAULT '',
                pages INTEGER DEFAULT 0,
                pdf_size INTEGER DEFAULT 0,
                pdf_size_formatted TEXT DEFAULT '',
                pdf_version TEXT DEFAULT '',
                error TEXT,
                created_at TEXT NOT NULL,
                finished_at TEXT
            )
        """)
        db.execute(
            "UPDATE jobs SET status='error', error='服务器重启，转换中断' WHERE status='processing'"
        )
        for col, default in [
            ("conversion_type", "'doc_to_pdf'"),
            ("docx_size", "0"),
            ("docx_size_formatted", "''"),
            ("markdown_size", "0"),
            ("markdown_size_formatted", "''"),
        ]:
            try:
                db.execute(f"ALTER TABLE jobs ADD COLUMN {col} TEXT NOT NULL DEFAULT {default}")
            except sqlite3.OperationalError:
                pass
        db.commit()


init_db()

# ── Word COM availability check ──
_word_available = None


def is_word_available():
    global _word_available
    if _word_available is None:
        try:
            import win32com.client  # noqa: F401
            _word_available = True
        except ImportError:
            _word_available = False
    return _word_available


# ── Conversion engine ──

def convert_doc_to_pdf(input_path: str, output_dir: str, options: dict | None = None) -> str:
    """Convert DOC/DOCX to PDF via Microsoft Word COM. Returns output PDF path."""
    import win32com.client

    word = None
    doc = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False

        abs_input = os.path.abspath(input_path)
        abs_output = os.path.join(
            output_dir,
            os.path.splitext(os.path.basename(input_path))[0] + ".pdf",
        )

        doc = word.Documents.Open(abs_input)

        # Apply page setup options if provided
        if options:
            _apply_page_setup(doc, options)
            _apply_font_settings(word, doc, options)
            _apply_image_quality_word(word, doc, options)

        # Determine export quality based on DPI setting and lossless flag
        image_dpi = parse_image_dpi(options)
        lossless = options.get("losslessImages", False) if options else False

        # For lossless mode, always use print quality optimization
        if lossless:
            optimize_for = 1  # wdExportOptimizeForPrint
        else:
            optimize_for = 1 if image_dpi >= 150 else 0  # 1=print, 0=screen

        doc.ExportAsFixedFormat(
            OutputFileName=abs_output,
            ExportFormat=17,           # wdExportFormatPDF
            OpenAfterExport=False,
            OptimizeFor=optimize_for,
            Range=0,                   # wdExportAllDocument
            Item=0,                    # wdExportDocumentContent
            IncludeDocProps=True,
            KeepIRM=True,
            CreateBookmarks=1,         # wdExportCreateWordBookmarks
            DocStructureTags=True,
            BitmapMissingFonts=True,
            UseISO19005_1=False,
        )

        if not os.path.isfile(abs_output):
            raise RuntimeError("Word conversion finished but output file not found")

        # Post-process PDF for additional options (password, permissions, PDF version)
        abs_output = _post_process_pdf(abs_output, options)

        return abs_output

    finally:
        if doc:
            try:
                doc.Close(SaveChanges=0)
            except Exception:
                pass
        if word:
            try:
                word.Quit()
            except Exception:
                pass


def _post_process_pdf(pdf_path: str, options: dict | None) -> str:
    """Post-process PDF for password protection, permissions, and PDF version.

    Uses pypdf to apply security settings and adjust PDF version.
    """
    if not options:
        return pdf_path

    # Check if any post-processing is needed
    password_protect = options.get("passwordProtect", False)
    pdf_version = options.get("pdfVersion", "1.7")
    allow_printing = options.get("allowPrinting", True)
    allow_copying = options.get("allowCopying", True)

    if not password_protect and pdf_version == "1.7" and allow_printing and allow_copying:
        return pdf_path

    try:
        from pypdf import PdfReader, PdfWriter
        from pypdf.constants import UserAccessFlags

        reader = PdfReader(pdf_path)
        writer = PdfWriter()

        # Copy all pages
        for page in reader.pages:
            writer.add_page(page)

        # Copy metadata
        if reader.metadata:
            writer.add_metadata(reader.metadata)

        # Set PDF version
        if pdf_version == "1.5":
            writer.pdf_version = "1.5"
        elif pdf_version == "1.6":
            writer.pdf_version = "1.6"
        else:
            writer.pdf_version = "1.7"

        # Calculate permission flags
        # Default: all permissions allowed
        permission_flags = (
            UserAccessFlags.PRINT |
            UserAccessFlags.MODIFY |
            UserAccessFlags.COPY |
            UserAccessFlags.ANNOTATE |
            UserAccessFlags.FILL_FORM |
            UserAccessFlags.EXTRACT |
            UserAccessFlags.ASSEMBLE |
            UserAccessFlags.PRINT_HQ
        )

        # Remove permissions based on settings
        if not allow_printing:
            permission_flags &= ~(UserAccessFlags.PRINT | UserAccessFlags.PRINT_HQ)
        if not allow_copying:
            permission_flags &= ~(UserAccessFlags.COPY | UserAccessFlags.EXTRACT)

        # Apply encryption with permissions
        if password_protect:
            password = options.get("password", "docflow")
            writer.encrypt(
                user_password=password,
                owner_password=password,
                permissions=permission_flags,
                algorithm="AES-256"
            )
        elif not allow_printing or not allow_copying:
            # Apply permissions even without password (using empty user password)
            writer.encrypt(
                user_password="",
                owner_password="docflow_owner",
                permissions=permission_flags,
                algorithm="AES-256"
            )

        # Write output
        temp_output = pdf_path + ".tmp"
        with open(temp_output, "wb") as f:
            writer.write(f)

        # Replace original file
        os.replace(temp_output, pdf_path)

    except Exception as e:
        logger.warning(f"PDF post-processing failed: {e}")

    return pdf_path


def _ocr_output_quality(docx_path: str) -> float:
    """Check OCR output quality by ratio of CJK characters to total text.

    Returns a value between 0 and 1. Low values indicate garbled/poor OCR.
    """
    try:
        from docx import Document
        doc = Document(docx_path)
        total_chars = 0
        cjk_chars = 0
        for para in doc.paragraphs:
            for run in para.runs:
                text = run.text
                total_chars += len(text)
                for ch in text:
                    cp = ord(ch)
                    if (0x4E00 <= cp <= 0x9FFF or  # CJK Unified
                        0x3400 <= cp <= 0x4DBF or  # CJK Extension A
                        0x3000 <= cp <= 0x303F or  # CJK Symbols
                        0xFF00 <= cp <= 0xFFEF):   # Fullwidth
                        cjk_chars += 1
        # Also check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            text = run.text
                            total_chars += len(text)
                            for ch in text:
                                cp = ord(ch)
                                if (0x4E00 <= cp <= 0x9FFF or
                                    0x3400 <= cp <= 0x4DBF or
                                    0x3000 <= cp <= 0x303F or
                                    0xFF00 <= cp <= 0xFFEF):
                                    cjk_chars += 1
        if total_chars == 0:
            return 0.0
        return cjk_chars / total_chars
    except Exception:
        return 0.0


def convert_pdf_to_docx(input_path: str, output_dir: str, options: dict | None = None) -> str:
    """Convert PDF to DOCX — routes scanned PDFs to OCR, native PDFs to pdf2docx.

    Supports mixed documents: per-page analysis routes native pages through
    pdf2docx and scanned pages through OCR, then merges results.
    """
    abs_input = os.path.abspath(input_path)
    abs_output = os.path.join(
        os.path.abspath(output_dir),
        os.path.splitext(os.path.basename(input_path))[0] + ".docx",
    )
    opts = options or {}
    use_ocr = opts.get("ocrEnabled", True)

    # --- Classify each page ---
    page_types = _classify_pages(abs_input)
    num_pages = len(page_types)
    num_scanned = sum(page_types)

    # --- All scanned: OCR path with structural fallback ---
    if use_ocr and num_scanned == num_pages:
        try:
            result = _ocr_scanned_pdf_to_docx(abs_input, abs_output, TEMP_DIR)
            # Check OCR quality: if output has very few recognizable characters,
            # fall back to structural extraction
            if _ocr_output_quality(abs_output) > 0.1:
                return result
        except Exception:
            pass
        # Fallback: structural form extraction
        return _structural_form_to_docx(abs_input, abs_output, TEMP_DIR)

    # --- All native: pdf2docx path (fast) ---
    if num_scanned == 0 or not use_ocr:
        return _convert_native_pdf(abs_input, abs_output, opts)

    # --- Mixed document: per-page routing ---
    return _convert_mixed_pdf(abs_input, abs_output, page_types, opts)


def _build_pdf2docx_kwargs(opts: dict) -> dict:
    """Build pdf2docx keyword arguments from user options."""
    conv_kwargs = {"ignore_page_error": True}

    # Image quality
    image_dpi = parse_image_dpi(opts)
    if opts.get("extractImages") is False:
        conv_kwargs["clip_image_res_ratio"] = 1.0
    else:
        conv_kwargs["clip_image_res_ratio"] = max(image_dpi / 72.0, 1.0)

    # Table processing
    table_mode = opts.get("tableMode", "自动")
    if table_mode == "严格":
        conv_kwargs.update({
            "parse_lattice_table": True,
            "parse_stream_table": True,
            "extract_stream_table": True,
            "connected_border_tolerance": 1.5,
            "max_border_width": 8.0,
            "min_border_clearance": 1.0,
            "shape_min_dimension": 1.0,
            "min_svg_gap_dx": 20.0,
            "min_svg_gap_dy": 5.0,
            "min_svg_w": 1.0,
            "min_svg_h": 1.0,
            "list_not_table": True,
        })
    elif table_mode == "宽松":
        conv_kwargs.update({
            "parse_lattice_table": True,
            "parse_stream_table": False,
            "extract_stream_table": False,
            "connected_border_tolerance": 0.3,
            "max_border_width": 4.0,
            "min_border_clearance": 3.0,
            "shape_min_dimension": 3.0,
            "list_not_table": True,
        })
    else:  # 自动
        conv_kwargs.update({
            "parse_lattice_table": True,
            "parse_stream_table": True,
            "extract_stream_table": False,
            "connected_border_tolerance": 1.0,
            "max_border_width": 6.0,
            "min_border_clearance": 1.5,
            "shape_min_dimension": 1.5,
            "min_svg_gap_dx": 18.0,
            "min_svg_gap_dy": 4.0,
            "min_svg_w": 1.5,
            "min_svg_h": 1.5,
            "list_not_table": True,
        })

    # Header/footer exclusion
    ignore_edges = opts.get("ignoreEdges", False)
    conv_kwargs["page_margin_factor_top"] = 0.05 if ignore_edges else 0.5
    conv_kwargs["page_margin_factor_bottom"] = 0.05 if ignore_edges else 0.5

    # Style fidelity
    keep_style = opts.get("keepStyle", True)
    if keep_style:
        conv_kwargs.update({
            "max_line_spacing_ratio": 2.0,
            "line_overlap_threshold": 0.95,
            "line_break_width_ratio": 0.4,
            "line_break_free_space_ratio": 0.15,
            "line_separate_threshold": 3.0,
            "new_paragraph_free_space_ratio": 0.8,
            "lines_left_aligned_threshold": 1.5,
            "lines_right_aligned_threshold": 1.5,
            "lines_center_aligned_threshold": 2.0,
            "delete_end_line_hyphen": True,
        })

    return conv_kwargs


def _convert_native_pdf(abs_input: str, abs_output: str, opts: dict) -> str:
    """Convert a fully native PDF using pdf2docx."""
    from pdf2docx import Converter
    conv_kwargs = _build_pdf2docx_kwargs(opts)

    cv = Converter(abs_input)
    try:
        cv.convert(abs_output, **conv_kwargs)
    finally:
        cv.close()

    if not os.path.isfile(abs_output):
        raise RuntimeError("PDF 转换完成但输出文件未找到")

    _remove_fullpage_images(abs_output)
    return abs_output


def _convert_mixed_pdf(abs_input: str, abs_output: str, page_types: list[bool], opts: dict) -> str:
    """Convert a mixed document with per-page routing.

    Groups consecutive pages of the same type, converts each group separately,
    then merges all results into a single DOCX.
    """
    from pdf2docx import Converter

    # Group consecutive pages by type: [(is_scanned, [page_indices]), ...]
    groups = []
    current_type = page_types[0]
    current_pages = [0]
    for i in range(1, len(page_types)):
        if page_types[i] == current_type:
            current_pages.append(i)
        else:
            groups.append((current_type, current_pages))
            current_type = page_types[i]
            current_pages = [i]
    groups.append((current_type, current_pages))

    temp_files = []
    conv_kwargs = _build_pdf2docx_kwargs(opts)

    try:
        for is_scanned, page_indices in groups:
            if is_scanned:
                # OCR these pages via a temporary single-page PDF
                import fitz
                src = fitz.open(abs_input)
                temp_pdf = os.path.join(TEMP_DIR, f"_mixed_{'_'.join(map(str, page_indices))}.pdf")
                dst = fitz.open()
                for idx in page_indices:
                    dst.insert_pdf(src, from_page=idx, to_page=idx)
                dst.save(temp_pdf)
                dst.close()
                src.close()

                temp_docx = os.path.join(TEMP_DIR, f"_mixed_{'_'.join(map(str, page_indices))}.docx")
                try:
                    _ocr_scanned_pdf_to_docx(temp_pdf, temp_docx, TEMP_DIR)
                    if _ocr_output_quality(temp_docx) < 0.1:
                        raise RuntimeError("OCR quality too low")
                except Exception:
                    _structural_form_to_docx(temp_pdf, temp_docx, TEMP_DIR)
                temp_files.append(temp_docx)

                # Cleanup temp PDF
                try:
                    os.remove(temp_pdf)
                except OSError:
                    pass
            else:
                # Convert native pages with pdf2docx
                cv = Converter(abs_input)
                temp_docx = os.path.join(TEMP_DIR, f"_mixed_native_{'_'.join(map(str, page_indices))}.docx")
                try:
                    if len(page_indices) == 1:
                        cv.convert(temp_docx, pages=page_indices, **conv_kwargs)
                    else:
                        # For consecutive pages, use start/end for multiprocessing support
                        cv.convert(temp_docx, start=page_indices[0], end=page_indices[-1], **conv_kwargs)
                finally:
                    cv.close()

                if os.path.isfile(temp_docx):
                    _remove_fullpage_images(temp_docx)
                    temp_files.append(temp_docx)

        # Merge all partial DOCX files
        if temp_files:
            _merge_docx_files(temp_files, abs_output)
        else:
            raise RuntimeError("混合转换未产生任何输出")

        if not os.path.isfile(abs_output):
            raise RuntimeError("PDF 转换完成但输出文件未找到")

        return abs_output

    finally:
        # Cleanup temp files
        for f in temp_files:
            try:
                os.remove(f)
            except OSError:
                pass


# ── Markdown converter (adapted from siyuan-tools) ──

_MD_LIST_RE = re.compile(
    r"^(\s*)"
    r"(?:"
    r"(\d{1,3}[.)、])"
    r"|(（\d{1,3}）)"
    r"|([•●○◆◇▪▸►‣∙⊙○◦])"
    r"|([-*+])"
    r")\s+"
)

_MD_SENTENCE_END_RE = re.compile(
    r'[。！？…；：）】」』》＞\s]$'
)

_MD_IMAGE_LABELS = frozenset({"figure", "picture"})
_MD_CODE_FENCE_RE = re.compile(r'^\s*```')
_MD_PAGE_SEP_RE = re.compile(r'\n*---\s*\n+<!--\s*第\s*\d+\s*页\s*-->\s*\n*')


@dataclass
class _LayoutRegion:
    x0: float
    y0: float
    x1: float
    y1: float
    label: str
    heading_level: int = 0


def _md_get_layout_regions(page) -> list:
    try:
        import pymupdf as _pymupdf
        raw = _pymupdf._get_layout(page)
    except (ValueError, Exception) as e:
        logger.debug("Markdown layout analysis failed: %s", e)
        return []
    regions = []
    for item in raw:
        if len(item) >= 5:
            regions.append(_LayoutRegion(
                x0=item[0], y0=item[1],
                x1=item[2], y1=item[3],
                label=item[4],
            ))
    return regions


def _md_assign_header_levels(page, regions: list) -> list:
    import pymupdf as _pymupdf
    headers = [r for r in regions if r.label == "section-header"]
    if not headers:
        return regions
    header_sizes = []
    for h in headers:
        rect = _pymupdf.Rect(h.x0, h.y0, h.x1, h.y1)
        try:
            blocks = page.get_text("dict", clip=rect, flags=_pymupdf.TEXT_PRESERVE_WHITESPACE)["blocks"]
        except ValueError:
            blocks = []
        max_size = 12.0
        for b in blocks:
            if b.get("type") != 0:
                continue
            for line in b.get("lines", []):
                for span in line.get("spans", []):
                    s = span.get("size", 12)
                    if s > max_size:
                        max_size = s
        header_sizes.append((h, max_size))
    unique_sizes = sorted(set(s for _, s in header_sizes), reverse=True)
    size_to_level = {s: i + 1 for i, s in enumerate(unique_sizes[:4])}
    for h, s in header_sizes:
        h.heading_level = size_to_level.get(s, 2)
    return regions


def _md_extract_region_text(page, region, pad: int = 6) -> str:
    import pymupdf as _pymupdf
    rect = _pymupdf.Rect(
        region.x0 - pad, region.y0 - pad,
        region.x1 + pad, region.y1 + pad,
    )
    try:
        blocks = page.get_text("dict", clip=rect, flags=_pymupdf.TEXT_PRESERVE_WHITESPACE)["blocks"]
    except ValueError:
        return ""
    lines = []
    for b in blocks:
        if b.get("type") != 0:
            continue
        for line_data in b.get("lines", []):
            parts = []
            for span in line_data.get("spans", []):
                t = span.get("text", "")
                if t:
                    parts.append(t)
            if parts:
                line_text = "".join(parts).strip()
                if line_text:
                    lines.append(line_text)
    if not lines:
        return ""
    merged = [lines[0]]
    for line in lines[1:]:
        prev = merged[-1]
        if _md_is_code_line(prev):
            merged.append(line)
            continue
        if (not _MD_SENTENCE_END_RE.search(prev)
            and not _MD_LIST_RE.match(line)
            and not re.match(r'^[\d]', line)
            and not line.startswith('#')
            and not _md_is_code_line(line)):
            merged[-1] += line
        else:
            merged.append(line)
    return "\n".join(merged).strip()


def _md_table_to_markdown(table_data) -> str:
    rows = []
    try:
        raw = table_data.extract()
    except ValueError:
        return ""
    for row in raw:
        cells = [str(c).replace("\n", " ").strip() if c is not None else "" for c in row]
        rows.append(cells)
    if not rows or not rows[0]:
        return ""
    col_count = max(len(r) for r in rows)
    rows = [r + [""] * (col_count - len(r)) for r in rows]
    lines = ["| " + " | ".join(rows[0]) + " |"]
    lines.append("| " + " | ".join(["---"] * col_count) + " |")
    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def _md_rect_overlap_area(r1: tuple, r2: tuple) -> float:
    x_overlap = max(0, min(r1[2], r2[2]) - max(r1[0], r2[0]))
    y_overlap = max(0, min(r1[3], r2[3]) - max(r1[1], r2[1]))
    return x_overlap * y_overlap


def _md_is_in_table(region, table_rects: list) -> bool:
    r = (region.x0, region.y0, region.x1, region.y1)
    area = max(1, (region.x1 - region.x0) * (region.y1 - region.y0))
    for tr in table_rects:
        overlap = _md_rect_overlap_area(r, tr)
        if overlap / area > 0.5:
            return True
    return False


def _md_find_page_images(page) -> list:
    try:
        raw = page.get_images(full=True)
    except (ValueError, RuntimeError):
        return []
    images = []
    for item in raw:
        if len(item) >= 4:
            images.append({"xref": item[0], "width": item[2], "height": item[3]})
    return images


def _md_extract_figure_image(doc, page, region, page_images, media_dir, page_idx, fig_idx):
    import pymupdf as _pymupdf
    try:
        fig_rect = (region.x0, region.y0, region.x1, region.y1)
        fig_area = max(1, (region.x1 - region.x0) * (region.y1 - region.y0))
        best_xref = None
        best_overlap = 0.0
        for img in page_images:
            try:
                rects = page.get_image_rects(img["xref"])
            except (ValueError, RuntimeError):
                continue
            for rect in rects:
                img_rect = (rect.x0, rect.y0, rect.x1, rect.y1)
                overlap = _md_rect_overlap_area(fig_rect, img_rect)
                ratio = overlap / fig_area
                if ratio > best_overlap:
                    best_overlap = ratio
                    best_xref = img["xref"]
        if best_xref is not None and best_overlap >= 0.3:
            img_data = doc.extract_image(best_xref)
            if img_data and img_data.get("image"):
                ext = img_data.get("ext", "png")
                filename = f"page{page_idx}_fig{fig_idx}.{ext}"
                filepath = os.path.join(media_dir, filename)
                with open(filepath, "wb") as f:
                    f.write(img_data["image"])
                return filename
        pad = 2
        rect = _pymupdf.Rect(
            region.x0 - pad, region.y0 - pad,
            region.x1 + pad, region.y1 + pad,
        )
        pixmap = page.get_pixmap(clip=rect, dpi=300)
        filename = f"page{page_idx}_fig{fig_idx}.png"
        filepath = os.path.join(media_dir, filename)
        pixmap.save(filepath)
        return filename
    except (ValueError, RuntimeError, OSError):
        return None


def _md_find_figure_caption(figure_region, all_regions, page, region_idx):
    if region_idx + 1 >= len(all_regions):
        return None, None
    next_region = all_regions[region_idx + 1]
    if next_region.label != "text":
        return None, None
    if next_region.y0 < figure_region.y1 - 5 or next_region.y0 > figure_region.y1 + 30:
        return None, None
    x_overlap = max(0, min(figure_region.x1, next_region.x1) - max(figure_region.x0, next_region.x0))
    fig_width = max(1, figure_region.x1 - figure_region.x0)
    if x_overlap / fig_width < 0.5:
        return None, None
    caption = _md_extract_region_text(page, next_region)
    if caption:
        return caption, region_idx + 1
    return None, None


def _md_is_code_line(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    if re.match(r'^\s*#\s*(include|define|ifdef|ifndef|endif|pragma|if|else|elif|undef)\b', stripped):
        return True
    if re.match(r'^\s*(void|int|char|float|double|struct|enum|union|typedef|static|const|extern|unsigned|long|short|signed|register|volatile)\b', stripped):
        return True
    if re.match(r'^\s*(if|else|for|while|do|switch|case|default|return|break|continue|goto)\b', stripped):
        return True
    if re.match(r'^\s*[{}]\s*$', stripped):
        return True
    if stripped.endswith(';') and len(stripped) > 3:
        return True
    if re.match(r'^\s*\w+.*,\s*$', stripped) and not re.search(r'[一-鿿]', stripped):
        return True
    if re.match(r'^\s*}\s*[;,]?\s*$', stripped):
        return True
    if re.match(r'^\s*(printf|scanf|malloc|free|sleep|getch|exit|main)\s*\(', stripped):
        return True
    if re.match(r'^\s*\w+\s*->\s*\w+', stripped):
        return True
    if (re.match(r'^[\s\w{}()\[\];,.<>!=&|+\-*/%^~?:\'"\\]+$', stripped)
            and len(stripped) > 2
            and re.search(r'[{}();,\[\]]', stripped)):
        return True
    return False


def _md_wrap_code_block(text: str) -> str:
    lines = text.split("\n")
    non_empty = [(i, l) for i, l in enumerate(lines) if l.strip()]
    if len(non_empty) < 3:
        return text
    code_count = sum(1 for _, l in non_empty if _md_is_code_line(l))
    if code_count / len(non_empty) < 0.3:
        return text
    code_start = 0
    for i, line in enumerate(lines):
        if _md_is_code_line(line):
            code_start = i
            break
    result_lines = []
    brace_depth = 0
    seen_opening_brace = False
    for line in lines[code_start:]:
        stripped = line.strip()
        if not stripped:
            result_lines.append(line)
            continue
        if brace_depth > 0:
            brace_depth += stripped.count('{') - stripped.count('}')
            result_lines.append(line)
        elif _md_is_code_line(line):
            result_lines.append(line)
            delta = stripped.count('{') - stripped.count('}')
            brace_depth = max(0, brace_depth + delta)
            if '{' in stripped:
                seen_opening_brace = True
            if brace_depth == 0 and stripped == '}' and not seen_opening_brace:
                brace_depth = 1
        else:
            break
    if len([l for l in result_lines if l.strip()]) < 3:
        return text
    return "```c\n" + "\n".join(result_lines).strip() + "\n```"


def _md_merge_paragraphs_raw(parts: list) -> list:
    merged = []
    buffer = ""
    list_buffer = []

    def flush_text():
        nonlocal buffer
        if buffer.strip():
            merged.append(buffer.strip())
        buffer = ""

    def flush_list():
        nonlocal list_buffer
        if list_buffer:
            merged.append("\n".join(list_buffer))
            list_buffer = []

    for part in parts:
        is_heading = part.startswith("#")
        is_list = bool(_MD_LIST_RE.match(part)) or part.startswith("1. ") or part.startswith("- ")
        is_code = _MD_CODE_FENCE_RE.match(part)
        if is_heading or is_code:
            flush_text()
            flush_list()
            merged.append(part)
            continue
        if is_list:
            flush_text()
            list_buffer.append(part)
            continue
        flush_list()
        if buffer:
            flush_text()
        buffer = part
    flush_text()
    flush_list()
    return merged


def _md_merge_cross_page_code(md: str) -> str:
    parts = re.split(r'(```[^\n]*\n.*?```)', md, flags=re.DOTALL)
    if len(parts) < 3:
        return md
    result = [parts[0]]
    i = 1
    while i < len(parts):
        part = parts[i]
        stripped_part = part.lstrip()
        if stripped_part.startswith('```'):
            code_block = stripped_part
            leading = part[:len(part) - len(stripped_part)]
            if leading:
                result.append(leading)
            lang_match = re.match(r'```([^\n]*)\n', code_block)
            lang = lang_match.group(1).strip() if lang_match else ''
            code_content = code_block[len(lang_match.group(0)):-3]
            j = i + 1
            while j + 1 < len(parts):
                gap = parts[j]
                stripped_gap = gap.strip()
                if stripped_gap == '' or _MD_PAGE_SEP_RE.fullmatch('\n' + stripped_gap + '\n'):
                    next_part = parts[j + 1].lstrip()
                    if next_part.startswith('```'):
                        next_block = next_part
                        next_lang_match = re.match(r'```([^\n]*)\n', next_block)
                        next_lang = next_lang_match.group(1).strip() if next_lang_match else ''
                        if lang == next_lang:
                            next_content = next_block[len(next_lang_match.group(0)):-3]
                            code_content = code_content + '\n' + next_content
                            j += 2
                            continue
                    j += 1
                    continue
                break
            merged_block = f'```{lang}\n{code_content}\n```'
            result.append(merged_block)
            i = j
        else:
            result.append(part)
            i += 1
    return ''.join(result)


def convert_to_markdown(input_path: str, output_dir: str, options: dict | None = None) -> tuple:
    """Convert a PDF to Markdown with image extraction.

    Returns (markdown_path, media_dir_path).
    """
    import pymupdf as _pymupdf
    try:
        from pymupdf import layout as _pymupdf_layout
        _pymupdf_layout.activate()
    except ImportError:
        pass

    abs_input = os.path.abspath(input_path)
    if not os.path.isfile(abs_input):
        raise FileNotFoundError(f"Input file not found: {abs_input}")

    base_name = os.path.splitext(os.path.basename(abs_input))[0]
    media_dir = os.path.join(os.path.abspath(output_dir), f"{base_name}_media")
    md_path = os.path.join(os.path.abspath(output_dir), base_name + ".md")

    all_sections = []
    doc = _pymupdf.open(abs_input)

    try:
        for i, page in enumerate(doc):
            page_parts = []

            # Step 1: Detect tables
            table_rects = []
            try:
                table_finder = page.find_tables()
            except ValueError:
                table_finder = None
            if table_finder and table_finder.tables:
                for table in table_finder.tables:
                    md = _md_table_to_markdown(table)
                    if md.strip():
                        page_parts.append(md)
                    try:
                        bbox = table.bbox
                        table_rects.append(bbox)
                    except (ValueError, RuntimeError):
                        pass

            # Step 2: AI layout regions
            regions = _md_get_layout_regions(page)
            if not regions:
                text = page.get_text("text").strip()
                if text:
                    page_parts.append(text)
                if page_parts:
                    if i > 0:
                        all_sections.append(f"\n\n---\n\n<!-- Page {i + 1} -->\n")
                    all_sections.append("\n\n".join(page_parts))
                continue

            regions = _md_assign_header_levels(page, regions)
            regions.sort(key=lambda r: (r.y0, r.x0))

            # Check for images
            has_figures = any(r.label in _MD_IMAGE_LABELS for r in regions)
            page_images = _md_find_page_images(page)
            if has_figures or page_images:
                os.makedirs(media_dir, exist_ok=True)

            figure_counter = 0
            covered_rects = [(r.x0, r.y0, r.x1, r.y1) for r in regions if r.label in _MD_IMAGE_LABELS]

            # Collect uncovered embedded images
            extra_images = []
            if page_images:
                for img in page_images:
                    try:
                        rects = page.get_image_rects(img["xref"])
                    except (ValueError, RuntimeError):
                        continue
                    for rect in rects:
                        img_rect = (rect.x0, rect.y0, rect.x1, rect.y1)
                        img_area = max(1, (rect.x1 - rect.x0) * (rect.y1 - rect.y0))
                        already_covered = any(
                            _md_rect_overlap_area(img_rect, cr) / img_area > 0.5
                            for cr in covered_rects
                        )
                        if already_covered:
                            continue
                        try:
                            img_data = doc.extract_image(img["xref"])
                            if img_data and img_data.get("image"):
                                ext = img_data.get("ext", "png")
                                filename = f"page{i}_fig{figure_counter}.{ext}"
                                filepath = os.path.join(media_dir, filename)
                                with open(filepath, "wb") as f:
                                    f.write(img_data["image"])
                                extra_images.append((rect.y0, filename, figure_counter))
                                figure_counter += 1
                        except (ValueError, RuntimeError, OSError):
                            pass

            # Step 3: Extract text from regions
            raw_parts = []
            consumed_indices = set()
            extra_idx = 0
            for region_idx, region in enumerate(regions):
                if region_idx in consumed_indices:
                    continue
                if region.label == "page-footer":
                    continue
                if _md_is_in_table(region, table_rects):
                    continue

                while extra_idx < len(extra_images) and extra_images[extra_idx][0] < region.y0:
                    _, fn, fig_num = extra_images[extra_idx]
                    raw_parts.append(("figure", f"![Figure {i + 1}-{fig_num + 1}]({fn})"))
                    extra_idx += 1

                if region.label in _MD_IMAGE_LABELS:
                    img_filename = _md_extract_figure_image(
                        doc, page, region, page_images, media_dir, i, figure_counter,
                    )
                    caption, cap_idx = _md_find_figure_caption(region, regions, page, region_idx)
                    if cap_idx is not None:
                        consumed_indices.add(cap_idx)
                    region_text = _md_extract_region_text(page, region)
                    if img_filename:
                        alt_text = caption or region_text or f"Figure {i + 1}-{figure_counter + 1}"
                        if len(alt_text) > 80:
                            alt_text = alt_text[:80] + "..."
                        raw_parts.append(("figure", f"![{alt_text}]({img_filename})"))
                    elif caption:
                        raw_parts.append(("text", caption))
                    if region_text:
                        raw_parts.append(("text", region_text))
                    figure_counter += 1
                    continue

                region_text = _md_extract_region_text(page, region)
                if not region_text:
                    continue
                raw_parts.append((region.label, region_text))

            while extra_idx < len(extra_images):
                _, fn, fig_num = extra_images[extra_idx]
                raw_parts.append(("figure", f"![Figure {i + 1}-{fig_num + 1}]({fn})"))
                extra_idx += 1

            # Step 4: Format as markdown
            text_parts = []
            for idx, (label, text) in enumerate(raw_parts):
                if label == "section-header":
                    level = 2
                    for r in regions:
                        if r.label == "section-header" and _md_extract_region_text(page, r) == text:
                            level = r.heading_level or 2
                            break
                    text_parts.append(f"{'#' * level} {text}")
                elif label == "list-item":
                    stripped = text.strip()
                    is_bare_marker = bool(re.match(r'^[\d]{1,3}[.)、]\s*$', stripped))
                    is_bare_bullet = bool(re.match(r'^[•●○◆◇▪▸►‣∙⊙○◦]\s*$', stripped))
                    if is_bare_marker or is_bare_bullet:
                        continue
                    if idx > 0:
                        prev_label, prev_text = raw_parts[idx - 1]
                        prev_stripped = prev_text.strip()
                        if prev_label == "list-item" and (
                            re.match(r'^[\d]{1,3}[.)、]\s*$', prev_stripped)
                            or re.match(r'^[•●○◆◇▪▸►‣∙⊙○◦]\s*$', prev_stripped)
                        ):
                            m = re.match(r'^[\d]+[.)、]', prev_stripped)
                            if m:
                                text_parts.append(f"1. {text}")
                            else:
                                text_parts.append(f"- {text}")
                            continue
                    lines = text.split("\n")
                    for line in lines:
                        line = line.strip()
                        if not line:
                            continue
                        m = _MD_LIST_RE.match(line)
                        if m:
                            content = line[m.end():]
                            if m.group(2) or m.group(3):
                                text_parts.append(f"1. {content}")
                            else:
                                text_parts.append(f"- {content}")
                        else:
                            text_parts.append(f"- {line}")
                elif label in ("text", "equation"):
                    text_parts.append(_md_wrap_code_block(text))
                elif label in _MD_IMAGE_LABELS:
                    text_parts.append(text)

            if text_parts:
                merged = _md_merge_paragraphs_raw(text_parts)
                page_parts.append("\n\n".join(merged))

            if page_parts:
                if i > 0:
                    all_sections.append(f"\n\n---\n\n<!-- Page {i + 1} -->\n")
                all_sections.append("\n\n".join(page_parts))
    finally:
        doc.close()

    markdown = _md_merge_cross_page_code("\n\n".join(all_sections))

    with open(md_path, "w", encoding="utf-8") as f:
        f.write(markdown)

    media_path = media_dir if os.path.isdir(media_dir) else None
    return md_path, media_path


def _find_tesseract() -> str | None:
    """Locate Tesseract executable. Returns path or None."""
    import shutil as _shutil
    tess_cmd = _shutil.which("tesseract")
    if tess_cmd:
        return tess_cmd
    for candidate in (
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    ):
        if os.path.isfile(candidate):
            return candidate
    return None


def _is_scanned_pdf(input_path: str) -> bool:
    """Detect if a PDF is predominantly scanned (image-only, no text layer)."""
    import fitz
    doc = fitz.open(input_path)
    scanned = 0
    total = len(doc)
    for page in doc:
        text = page.get_text("text").strip()
        images = page.get_images(full=True)
        if not images:
            continue
        # Check ALL images on the page, find the largest one
        max_coverage = 0.0
        for img_info in images:
            img_rects = page.get_image_rects(img_info[0])
            for r in img_rects:
                page_rect = page.rect
                coverage = (r.width * r.height) / (page_rect.width * page_rect.height)
                if coverage > max_coverage:
                    max_coverage = coverage
        # Page is "scanned" if it has very little text AND a dominant image
        if len(text) < 100 and max_coverage > 0.6:
            scanned += 1
    doc.close()
    return scanned > total * 0.3 if total > 0 else False


def _classify_pages(input_path: str) -> list[bool]:
    """Classify each page as scanned (True) or native (False).

    Returns a list of booleans, one per page.
    """
    import fitz
    doc = fitz.open(input_path)
    result = []
    for page in doc:
        text = page.get_text("text").strip()
        images = page.get_images(full=True)
        is_scanned = False
        if images:
            max_coverage = 0.0
            for img_info in images:
                img_rects = page.get_image_rects(img_info[0])
                for r in img_rects:
                    page_rect = page.rect
                    coverage = (r.width * r.height) / (page_rect.width * page_rect.height)
                    if coverage > max_coverage:
                        max_coverage = coverage
            if len(text) < 100 and max_coverage > 0.6:
                is_scanned = True
        result.append(is_scanned)
    doc.close()
    return result


def _merge_docx_files(docx_paths: list[str], output_path: str) -> str:
    """Merge multiple DOCX files into one by appending content from each.

    The first file's styles/sections are used as the base.
    """
    from docx import Document

    if not docx_paths:
        raise RuntimeError("没有可合并的 DOCX 文件")

    if len(docx_paths) == 1:
        import shutil
        shutil.copy2(docx_paths[0], output_path)
        return output_path

    base = Document(docx_paths[0])

    for path in docx_paths[1:]:
        src = Document(path)
        for element in src.element.body:
            base.element.body.append(element)

    base.save(output_path)
    return output_path


def _preprocess_scanned_image(img_path: str) -> str:
    """Preprocess a scanned page image for better OCR accuracy.

    Applies grayscale conversion, contrast enhancement (CLAHE), and
    adaptive thresholding. Returns the path to the preprocessed image.
    """
    try:
        import cv2
        import numpy as np

        img = cv2.imread(img_path)
        if img is None:
            return img_path

        # Convert to grayscale
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

        # Denoise
        denoised = cv2.fastNlMeansDenoising(gray, h=10)

        # CLAHE contrast enhancement
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        enhanced = clahe.apply(denoised)

        # Adaptive thresholding for clean binarization
        binary = cv2.adaptiveThreshold(
            enhanced, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY, 15, 8
        )

        # Save preprocessed image (overwrite original)
        cv2.imwrite(img_path, binary)
        return img_path
    except Exception:
        return img_path


def _ocr_scan_page(img_path: str, tess_cmd: str = None,
                    dpi: int = 300) -> list[tuple[str, float, float, float, float]]:
    """OCR a single page image, return list of (text, x, y, w, h) in PDF points.

    优先使用 RapidOCR（PaddleOCR ONNX），不可用时 fallback 到 Tesseract。
    """
    # --- RapidOCR (primary engine) ---
    engine = _get_ocr_engine()
    if engine is not None:
        try:
            words = engine.ocr_page(img_path, dpi=dpi)
            if words:
                return words
        except Exception as e:
            logger.warning(f"RapidOCR 失败，降级到 Tesseract: {e}")

    # --- Tesseract (fallback) ---
    if tess_cmd is None:
        tess_cmd = _find_tesseract()
    if not tess_cmd:
        return []

    import subprocess, csv
    tsv_path = os.path.splitext(img_path)[0] + ".tsv"
    subprocess.run(
        [tess_cmd, img_path, os.path.splitext(tsv_path)[0],
         "--dpi", str(dpi), "-l", "chi_sim+eng", "tsv"],
        capture_output=True, timeout=120,
    )
    words = []
    if not os.path.isfile(tsv_path):
        return words
    scale = 72.0 / dpi
    with open(tsv_path, "r", encoding="utf-8") as f:
        for row in csv.DictReader(f, delimiter="\t"):
            try:
                conf = float(row.get("conf", "-1"))
                text_val = row.get("text", "").strip()
                if conf > 50 and text_val:
                    x = int(row["left"]) * scale
                    y = int(row["top"]) * scale
                    w = int(row["width"]) * scale
                    h = int(row["height"]) * scale
                    words.append((text_val, x, y, w, h))
            except (KeyError, ValueError):
                continue
    try:
        os.remove(tsv_path)
    except OSError:
        pass
    return words


def _ocr_scanned_pdf_to_docx(input_path: str, output_path: str, temp_dir: str) -> str:
    """Convert a scanned PDF to DOCX via OCR with format preservation.

    Detects table structures from word positions, infers font sizes from bounding
    box heights, and recreates the original layout as closely as possible.
    优先使用 RapidOCR，不可用时降级到 Tesseract。
    """
    import fitz
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.oxml.ns import qn

    # Check if any OCR engine is available
    engine = _get_ocr_engine()
    tess_cmd = _find_tesseract()
    if engine is None and not tess_cmd:
        raise RuntimeError(
            "未安装 OCR 引擎，无法转换扫描版 PDF。"
            "请安装: pip install rapidocr-onnxruntime"
        )

    src = fitz.open(input_path)
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    for i, page in enumerate(src):
        page_rect = page.rect
        width_in = page_rect.width / 72.0
        height_in = page_rect.height / 72.0

        if i == 0:
            section = doc.sections[0]
        else:
            section = doc.add_section()
        section.page_width = Inches(width_in)
        section.page_height = Inches(height_in)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

        pix = page.get_pixmap(dpi=300)
        img_path = os.path.join(temp_dir, f"_ocr_p{i}.png")
        pix.save(img_path)

        # Preprocess scanned image for better OCR accuracy
        _preprocess_scanned_image(img_path)

        words = _ocr_scan_page(img_path, tess_cmd)

        if words:
            table_data = _detect_table_grid(words, page_rect.width, page_rect.height)

            if table_data:
                _add_ocr_table(doc, table_data, words, width_in)
            else:
                _add_ocr_paragraphs(doc, words)
        else:
            para = doc.add_paragraph("（此页无可识别文字）")

        try:
            os.remove(img_path)
        except OSError:
            pass

    src.close()
    doc.save(output_path)
    return output_path


def _detect_table_grid(
    words: list[tuple[str, float, float, float, float]],
    page_w: float,
    page_h: float,
) -> dict | None:
    """Detect table grid from OCR word positions.

    Returns dict with 'rows', 'cols', 'grid', 'col_widths' or None if no table.
    grid[row][col] = list of (text, x, y, w, h) tuples in that cell.
    """
    if len(words) < 6:
        return None

    margin_x = page_w * 0.05
    margin_y = page_h * 0.05

    # --- Find horizontal gaps (row separators) ---
    by_y = sorted(words, key=lambda t: t[2] + t[4] / 2)
    y_mids = [t[2] + t[4] / 2 for t in by_y]
    y_gaps = []
    for j in range(1, len(y_mids)):
        gap = y_mids[j] - y_mids[j - 1]
        if gap > 4:
            y_gaps.append((gap, (y_mids[j - 1] + y_mids[j]) / 2))

    y_gaps.sort(reverse=True)
    row_seps = [0.0, page_h]
    for gap, mid in y_gaps:
        if len(row_seps) >= 20:
            break
        if all(abs(mid - s) > 12 for s in row_seps):
            row_seps.append(mid)
    row_seps.sort()

    # --- Find vertical gaps (column separators) ---
    by_x = sorted(words, key=lambda t: t[1])
    x_starts = [t[1] for t in by_x]
    x_gaps = []
    for j in range(1, len(x_starts)):
        gap = x_starts[j] - (by_x[j - 1][1] + by_x[j - 1][3])
        if gap > 8:
            x_gaps.append((gap, (by_x[j - 1][1] + by_x[j - 1][3] + x_starts[j]) / 2))

    x_gaps.sort(reverse=True)
    col_seps = [0.0, page_w]
    for gap, mid in x_gaps:
        if len(col_seps) >= 15:
            break
        if all(abs(mid - s) > 15 for s in col_seps):
            col_seps.append(mid)
    col_seps.sort()

    # Filter narrow columns
    col_seps = [col_seps[0]] + [
        col_seps[j]
        for j in range(1, len(col_seps) - 1)
        if col_seps[j + 1] - col_seps[j] > 20
    ] + [col_seps[-1]]

    if len(row_seps) < 3 or len(col_seps) < 3:
        return None

    # --- Map words to grid cells ---
    num_rows = len(row_seps) - 1
    num_cols = len(col_seps) - 1
    grid: list[list[list]] = [[[] for _ in range(num_cols)] for _ in range(num_rows)]

    for t in words:
        text_val, x, y, w, h = t
        y_mid = y + h / 2
        x_mid = x + w / 2
        if y_mid < margin_y or y_mid > page_h - margin_y:
            continue
        if x_mid < margin_x or x_mid > page_w - margin_x:
            continue

        ri = -1
        for r in range(num_rows):
            if row_seps[r] <= y_mid <= row_seps[r + 1]:
                ri = r
                break
        if ri < 0:
            continue

        ci = -1
        for c in range(num_cols):
            if col_seps[c] <= x_mid <= col_seps[c + 1]:
                ci = c
                break
        if ci < 0:
            continue

        grid[ri][ci].append(t)

    # Require at least 40% cell occupancy
    occupied = sum(1 for r in range(num_rows) for c in range(num_cols) if grid[r][c])
    total = num_rows * num_cols
    if occupied < total * 0.4:
        return None

    # Compute column widths
    col_widths = [
        (col_seps[c + 1] - col_seps[c]) / 72.0 for c in range(num_cols)
    ]

    return {
        "rows": row_seps,
        "cols": col_seps,
        "grid": grid,
        "col_widths": col_widths,
    }


def _add_ocr_table(doc, table_data: dict, all_words: list, width_in: float):
    """Add a formatted table to the document from detected grid data."""
    from docx.shared import Pt, Emu
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    grid = table_data["grid"]
    col_widths = table_data["col_widths"]
    num_rows = len(grid)
    num_cols = len(grid[0]) if grid else 0

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
        borders.append(element)
    tbl_pr.append(borders)

    # Set column widths
    for c, w in enumerate(col_widths):
        for row in table.rows:
            row.cells[c].width = Emu(int(w * 914400))

    body_size = _infer_font_size(all_words)

    for ri in range(num_rows):
        for ci in range(num_cols):
            cell = table.rows[ri].cells[ci]
            cell_words = grid[ri][ci]

            # Clear default empty paragraph
            for p in cell.paragraphs:
                if not p.text.strip():
                    p_elem = p._element
                    p_elem.getparent().remove(p_elem)

            if not cell_words:
                para = cell.add_paragraph()
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after = Pt(2)
                run = para.add_run(" ")
                run.font.size = Pt(body_size)
                run.font.name = "宋体"
                run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                continue

            cell_words_sorted = sorted(cell_words, key=lambda t: (t[2], t[1]))

            # Detect alignment
            cell_w = (table_data["cols"][ci + 1] - table_data["cols"][ci]) / 72.0
            word_positions = [t[1] / 72.0 for t in cell_words_sorted]
            avg_pos = sum(word_positions) / len(word_positions) if word_positions else 0
            avg_end = sum((t[1] + t[3]) / 72.0 for t in cell_words_sorted) / len(cell_words_sorted)
            cell_left = table_data["cols"][ci] / 72.0
            cell_right = table_data["cols"][ci + 1] / 72.0

            if avg_pos - cell_left < cell_w * 0.15 and avg_end - cell_left < cell_w * 0.5:
                alignment = 0  # LEFT
            elif cell_right - avg_end < cell_w * 0.15 and cell_right - avg_pos < cell_w * 0.5:
                alignment = 2  # RIGHT
            else:
                alignment = 1  # CENTER

            # Group words into sub-lines within the cell
            sub_lines: dict[int, list] = {}
            for t in cell_words_sorted:
                lk = round(t[2] / 8) * 8
                sub_lines.setdefault(lk, []).append(t)

            for lk in sorted(sub_lines.keys()):
                sub_words = sorted(sub_lines[lk], key=lambda t: t[1])
                # Determine font size for this sub-line
                heights = [t[4] for t in sub_words]
                median_h = sorted(heights)[len(heights) // 2]
                fs = max(8.0, min(36.0, round(median_h / 2.8 * 2) / 2))

                line_text = _join_words_with_spaces(sub_words)
                if not line_text.strip():
                    continue

                para = cell.add_paragraph(line_text)
                para.alignment = alignment
                para.paragraph_format.space_before = Pt(1)
                para.paragraph_format.space_after = Pt(1)
                para.paragraph_format.line_spacing = Pt(fs * 1.3)

                for run in para.runs:
                    run.font.size = Pt(fs)
                    run.font.name = "宋体"
                    run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _add_ocr_paragraphs(doc, words: list):
    """Add OCR text as paragraphs (no table structure detected)."""
    from docx.shared import Pt
    from docx.oxml.ns import qn

    body_size = _infer_font_size(words)

    lines: dict[int, list] = {}
    for t in words:
        lk = round(t[2] / 10) * 10
        lines.setdefault(lk, []).append(t)

    for lk in sorted(lines.keys()):
        line_words = sorted(lines[lk], key=lambda t: t[1])
        line_text = _join_words_with_spaces(line_words)
        if not line_text.strip():
            continue

        # Infer font size from this line's word heights
        heights = [t[4] for t in line_words]
        median_h = sorted(heights)[len(heights) // 2]
        fs = max(8.0, min(36.0, round(median_h / 2.8 * 2) / 2))

        # Detect alignment
        page_width_pt = max(t[1] + t[3] for t in words) if words else 595
        left_margin = min(t[1] for t in words)
        right_edge = max(t[1] + t[3] for t in words)
        line_left = line_words[0][1]
        line_right = line_words[-1][1] + line_words[-1][3]

        if abs(line_left - left_margin) < 15 and abs(line_right - right_edge) < 15:
            alignment = 1  # CENTER
        elif abs(line_left - left_margin) < 10:
            alignment = 0  # LEFT
        elif abs(line_right - right_edge) < 10:
            alignment = 2  # RIGHT
        else:
            alignment = 0

        para = doc.add_paragraph(line_text)
        para.alignment = alignment
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)

        for run in para.runs:
            run.font.size = Pt(fs)
            run.font.name = "宋体"
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _join_words_with_spaces(words: list[tuple[str, float, float, float, float]]) -> str:
    """Join OCR words with intelligently inserted spaces based on bounding box gaps.

    Compares the gap between consecutive words to the average character width
    and inserts a space when the gap exceeds ~0.3x the average char width.
    """
    if not words:
        return ""
    if len(words) == 1:
        return words[0][0]

    # Estimate average character width from all words
    total_chars = sum(len(t[0]) for t in words)
    total_width = sum(t[3] for t in words)
    avg_char_w = total_width / max(total_chars, 1)

    parts = [words[0][0]]
    for prev, cur in zip(words, words[1:]):
        gap = cur[1] - (prev[1] + prev[3])  # gap between prev end and cur start
        if gap > avg_char_w * 0.3:
            parts.append(" ")
        parts.append(cur[0])
    return "".join(parts)


def _infer_font_size(words: list) -> float:
    """Infer the dominant body text font size from word bounding box heights."""
    if not words:
        return 11.0
    heights = sorted(t[4] for t in words)
    median_h = heights[len(heights) // 2]
    size = median_h / 2.8
    return max(8.0, min(36.0, round(size * 2) / 2))


def _is_fullpage_image(block) -> bool:
    """Check if a DOCX block is a full-page placeholder image.

    These are typically rasterized scanned pages embedded by pdf2docx.
    Returns True if the image covers >75% of a standard page.
    """
    try:
        from docx.oxml.ns import qn
        inline = block._element.findall(qn("w:r/w:drawing/wp:inline"))
        if not inline:
            inline = block._element.findall(
                qn("w:r/w:drawing/wp:anchor")
            )
        for drawing in inline:
            extent = drawing.find(qn("wp:extent"))
            if extent is not None:
                cx = int(extent.get("cx", 0))
                cy = int(extent.get("cy", 0))
                # EMU to inches: divide by 914400
                w_in = cx / 914400
                h_in = cy / 914400
                area = w_in * h_in
                # Use relative threshold: image must be >75% of a typical page
                # and cover most of both dimensions (portrait or landscape)
                min_page_area = 70  # ~A5 minimum
                if area > min_page_area and w_in > 6.0 and h_in > 8.0:
                    return True
    except Exception:
        pass
    return False


def _remove_fullpage_images(docx_path: str) -> str:
    """Post-process DOCX: remove full-page placeholder images.

    These are typically scanned page rasterizations that pdf2docx embeds.
    After removal, only editable text and legitimate smaller images remain.
    Checks both paragraphs and table cells.
    """
    try:
        from docx import Document
        doc = Document(docx_path)
        removed = 0

        # Check paragraphs
        for para in list(doc.paragraphs):
            if _is_fullpage_image(para):
                p_elem = para._element
                p_elem.getparent().remove(p_elem)
                removed += 1

        # Check table cells — pdf2docx sometimes places placeholder images inside cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in list(cell.paragraphs):
                        if _is_fullpage_image(para):
                            p_elem = para._element
                            p_elem.getparent().remove(p_elem)
                            removed += 1

        if removed:
            doc.save(docx_path)
        return docx_path
    except Exception:
        return docx_path


def _detect_form_grid(image_path: str) -> dict | None:
    """Detect table grid structure from a form image using projection profiles.

    Uses horizontal projection for row separators and per-row gap analysis
    for column separators (more robust for JPEG-compressed scans).

    Returns dict with 'h_lines', 'v_lines' (pixel positions), 'image_size' (w,h),
    or None if no grid is detected.
    """
    import cv2
    import numpy as np

    img = cv2.imread(image_path)
    if img is None:
        return None
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    h, w = gray.shape

    _, binary = cv2.threshold(gray, 200, 255, cv2.THRESH_BINARY_INV)

    # --- Horizontal lines via projection ---
    h_proj = np.sum(binary, axis=1) / 255
    h_threshold = w * 0.2
    h_rows = np.where(h_proj > h_threshold)[0]

    def _cluster_lines(rows, gap=8):
        if len(rows) == 0:
            return []
        rows = sorted(rows)
        clusters = [[rows[0]]]
        for r in rows[1:]:
            if r - clusters[-1][-1] <= gap:
                clusters[-1].append(r)
            else:
                clusters.append([r])
        return sorted([int(np.mean(c)) for c in clusters])

    h_lines = _cluster_lines(h_rows, gap=8)

    # Filter horizontal lines: must span >20% of width
    h_filtered = []
    for y in h_lines:
        y_slice = binary[max(0, y-2):y+3, :]
        dark_ratio = np.sum(y_slice > 0) / (w * min(5, h))
        if dark_ratio > 0.15:
            h_filtered.append(y)

    # Merge very close horizontal lines (< 3% of image height apart)
    min_row_h = h * 0.03
    h_merged = [h_filtered[0]] if h_filtered else []
    for y in h_filtered[1:]:
        if y - h_merged[-1] < min_row_h:
            h_merged[-1] = (h_merged[-1] + y) // 2  # Average position
        else:
            h_merged.append(y)
    h_filtered = h_merged

    if len(h_filtered) < 2:
        return None

    # --- Vertical lines via per-row gap analysis ---
    # For each row region, find columns with very low projection (cell gaps)
    all_gap_positions = []
    for i in range(len(h_filtered) - 1):
        y1 = h_filtered[i]
        y2 = h_filtered[i + 1]
        if y2 - y1 < 30:
            continue

        row_binary = binary[y1:y2, :]
        v_proj = np.sum(row_binary, axis=0) / 255

        # Smooth
        kernel_size = max(5, w // 100)
        if kernel_size % 2 == 0:
            kernel_size += 1
        from scipy.ndimage import uniform_filter1d
        smoothed = uniform_filter1d(v_proj.astype(float), size=kernel_size)

        # Find low-projection gaps (cell separators)
        mean_proj = smoothed.mean()
        gap_thresh = mean_proj * 0.15

        in_gap = False
        gap_start = 0
        for x in range(w):
            if smoothed[x] < gap_thresh and not in_gap:
                gap_start = x
                in_gap = True
            elif smoothed[x] >= gap_thresh and in_gap:
                if x - gap_start > 10:
                    all_gap_positions.append((gap_start + x) // 2)
                in_gap = False

    # Cluster all gap positions to find consistent column separators
    if all_gap_positions:
        v_lines = _cluster_lines(all_gap_positions, gap=25)
        # Filter: keep gaps that appear in multiple rows
        from collections import Counter
        gap_counts = Counter()
        for pos in all_gap_positions:
            # Find which cluster this belongs to
            for vl in v_lines:
                if abs(pos - vl) < 25:
                    gap_counts[vl] += 1
                    break
        # Keep gaps that appear in at least 1 row (since some rows have few cells)
        v_filtered = sorted([x for x, cnt in gap_counts.items() if cnt >= 1])
    else:
        v_filtered = []

    # Add outer borders if not present
    if not v_filtered or v_filtered[0] > 20:
        v_filtered.insert(0, 0)
    if not v_filtered or v_filtered[-1] < w - 20:
        v_filtered.append(w)

    # Deduplicate
    v_filtered = sorted(set(v_filtered))

    # Merge narrow columns (< 5% of image width) with neighbors
    # These are typically text gaps, not real cell separators
    min_col_width = w * 0.05
    merged = [v_filtered[0]]
    for i in range(1, len(v_filtered)):
        col_w = v_filtered[i] - merged[-1]
        if col_w < min_col_width and i < len(v_filtered) - 1:
            continue  # Skip this separator, merge with next
        merged.append(v_filtered[i])
    v_filtered = merged

    if len(h_filtered) < 2 or len(v_filtered) < 2:
        return None

    return {
        "h_lines": sorted(h_filtered),
        "v_lines": sorted(v_filtered),
        "image_size": (w, h),
    }


def _ocr_cell_region(image_path: str, x1: int, y1: int, x2: int, y2: int,
                      tess_cmd: str = None) -> str:
    """OCR a single cell region and return recognized text."""
    import cv2

    img = cv2.imread(image_path)
    if img is None:
        return ""

    # Add padding
    h, w = img.shape[:2]
    pad = 5
    y1p = max(0, y1 - pad)
    y2p = min(h, y2 + pad)
    x1p = max(0, x1 - pad)
    x2p = min(w, x2 + pad)

    crop = img[y1p:y2p, x1p:x2p]
    if crop.size == 0:
        return ""

    # Preprocess: sharpen + threshold
    gray = cv2.cvtColor(crop, cv2.COLOR_BGR2GRAY)
    clahe = cv2.createCLAHE(clipLimit=4.0, tileGridSize=(8, 8))
    enhanced = clahe.apply(gray)
    _, binary = cv2.threshold(enhanced, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

    cell_path = image_path + f"_cell_{x1}_{y1}.png"
    cv2.imwrite(cell_path, binary)

    # Try RapidOCR first
    engine = _get_ocr_engine()
    if engine is not None:
        try:
            regions = engine.recognize(cell_path)
            if regions:
                text = " ".join(r["text"] for r in regions)
                return text
        except Exception as e:
            logger.warning(f"RapidOCR cell 失败: {e}")
        finally:
            try:
                os.remove(cell_path)
            except OSError:
                pass
        return ""

    # Fallback to Tesseract
    if tess_cmd is None:
        tess_cmd = _find_tesseract()
    if not tess_cmd:
        try:
            os.remove(cell_path)
        except OSError:
            pass
        return ""

    import subprocess
    try:
        result = subprocess.run(
            [tess_cmd, cell_path, cell_path + "_ocr",
             "--dpi", "300", "-l", "chi_sim", "--psm", "7"],
            capture_output=True, text=True, timeout=30,
        )
        txt_path = cell_path + "_ocr.txt"
        if os.path.isfile(txt_path):
            with open(txt_path, "r", encoding="utf-8") as f:
                text = f.read().strip()
            return text
    except Exception:
        pass
    finally:
        for suffix in ("", "_ocr.txt", "_ocr.tsv"):
            try:
                os.remove(cell_path + suffix)
            except OSError:
                pass

    return ""


def _reconstruct_form_docx(grid: dict, image_path: str, output_path: str) -> str:
    """Create a DOCX from a detected form grid structure.

    Crops each cell, attempts OCR, and builds a table with the recognized text.
    Cells that fail OCR get the original image embedded as reference.
    """
    from docx import Document
    from docx.shared import Inches, Pt, Emu
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    h_lines = grid["h_lines"]
    v_lines = grid["v_lines"]
    img_w, img_h = grid["image_size"]

    num_rows = len(h_lines) - 1
    num_cols = len(v_lines) - 1

    if num_rows < 1 or num_cols < 1:
        raise RuntimeError("表格结构检测失败")

    # Page dimensions (match source)
    page_w_in = img_w / 300.0  # assuming 300 DPI source
    page_h_in = img_h / 300.0

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(page_w_in)
    section.page_height = Inches(page_h_in)
    section.top_margin = Inches(0.3)
    section.bottom_margin = Inches(0.3)
    section.left_margin = Inches(0.3)
    section.right_margin = Inches(0.3)

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # Create table
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Add borders
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
        borders.append(element)
    tbl_pr.append(borders)

    # Set column widths
    for ci in range(num_cols):
        col_w_in = (v_lines[ci + 1] - v_lines[ci]) / 300.0
        for row in table.rows:
            row.cells[ci].width = Emu(int(col_w_in * 914400))

    # Fill cells
    for ri in range(num_rows):
        for ci in range(num_cols):
            cell = table.rows[ri].cells[ci]

            # Crop cell region from image (pixel coords)
            x1 = v_lines[ci]
            x2 = v_lines[ci + 1]
            y1 = h_lines[ri]
            y2 = h_lines[ri + 1]

            # Cell dimensions
            cell_w_px = x2 - x1
            cell_h_px = y2 - y1

            # Skip very small cells
            if cell_w_px < 20 or cell_h_px < 20:
                continue

            # Try OCR on this cell
            text = _ocr_cell_region(image_path, x1, y1, x2, y2)

            # Clear default paragraph
            for p in cell.paragraphs:
                if not p.text.strip():
                    p_elem = p._element
                    p_elem.getparent().remove(p_elem)

            if text.strip():
                # OCR succeeded — add text
                para = cell.add_paragraph(text.strip())
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after = Pt(2)
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.name = "宋体"
                    run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            else:
                # OCR failed — embed cell image as reference
                import cv2
                img = cv2.imread(image_path)
                if img is not None:
                    crop = img[max(0, y1):min(img.shape[0], y2),
                               max(0, x1):min(img.shape[1], x2)]
                    if crop.size > 0:
                        cell_img_path = image_path + f"_cellimg_{ri}_{ci}.png"
                        cv2.imwrite(cell_img_path, crop)
                        try:
                            para = cell.add_paragraph()
                            para.alignment = 1  # CENTER
                            run = para.add_run()
                            # Scale to fit cell width
                            cell_w_in = cell_w_px / 300.0
                            run.add_picture(cell_img_path, width=Inches(min(cell_w_in, 6.0)))
                        except Exception:
                            para = cell.add_paragraph(" ")
                        finally:
                            try:
                                os.remove(cell_img_path)
                            except OSError:
                                pass
                    else:
                        cell.add_paragraph(" ")
                else:
                    cell.add_paragraph(" ")

    doc.save(output_path)
    return output_path


def _structural_form_to_docx(input_path: str, output_path: str,
                              temp_dir: str) -> str:
    """Extract form structure from a scanned/image PDF and reconstruct as DOCX.

    This is the fallback when standard OCR fails (e.g., heavily compressed scans).
    Detects the table grid, crops each cell, attempts per-cell OCR, and builds
    a structured DOCX with the detected layout.
    """
    import fitz

    engine = _get_ocr_engine()
    tess_cmd = _find_tesseract()
    if engine is None and not tess_cmd:
        raise RuntimeError(
            "未安装 OCR 引擎，无法转换扫描版 PDF。"
            "请安装 rapidocr-onnxruntime 或 Tesseract。"
        )

    # Extract the page image
    src = fitz.open(input_path)
    page = src[0]
    pix = page.get_pixmap(dpi=300)
    img_path = os.path.join(temp_dir, "_structural_form.png")
    pix.save(img_path)
    src.close()

    # Detect grid
    grid = _detect_form_grid(img_path)
    if not grid:
        # Fallback: just embed the full page image
        from docx import Document
        from docx.shared import Inches
        doc = Document()
        section = doc.sections[0]
        page_rect = fitz.open(input_path)[0].rect
        section.page_width = Inches(page_rect.width / 72.0)
        section.page_height = Inches(page_rect.height / 72.0)
        para = doc.add_paragraph()
        para.alignment = 1
        run = para.add_run()
        run.add_picture(img_path, width=Inches(section.page_width.inches * 0.9))
        doc.save(output_path)
        return output_path

    return _reconstruct_form_docx(grid, img_path, output_path)


def get_docx_metadata(docx_path: str) -> dict:
    """Get basic metadata for a DOCX file."""
    file_size = os.path.getsize(docx_path) if os.path.isfile(docx_path) else 0
    return {"size": file_size, "sizeFormatted": format_size(file_size)}


def analyze_docx_images(docx_path: str) -> dict:
    """Analyze images in a DOCX file to determine optimal export settings.

    Returns info about image count, max resolution, and recommended DPI.
    """
    try:
        import zipfile
        import struct

        def parse_jpeg_dimensions(data: bytes) -> tuple[int, int]:
            """Parse JPEG dimensions from image data."""
            try:
                i = 0
                if data[0:2] != b'\xff\xd8':  # Not a JPEG
                    return 0, 0

                i = 2
                while i < len(data) - 1:
                    if data[i] != 0xFF:
                        i += 1
                        continue

                    marker = data[i + 1]
                    if marker == 0x00:
                        i += 2
                        continue

                    if marker in (0xC0, 0xC1, 0xC2):  # SOF0, SOF1, SOF2
                        if i + 9 < len(data):
                            height = struct.unpack('>H', data[i + 5:i + 7])[0]
                            width = struct.unpack('>H', data[i + 7:i + 9])[0]
                            return width, height

                    if i + 3 < len(data):
                        length = struct.unpack('>H', data[i + 2:i + 4])[0]
                        i += 2 + length
                    else:
                        break
                return 0, 0
            except Exception:
                return 0, 0

        image_count = 0
        max_width = 0
        max_height = 0
        total_image_size = 0

        with zipfile.ZipFile(docx_path, 'r') as docx_zip:
            # Look for images in the media folder
            for file_name in docx_zip.namelist():
                if file_name.startswith('word/media/') and file_name.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff')):
                    image_count += 1
                    file_info = docx_zip.getinfo(file_name)
                    total_image_size += file_info.file_size

                    # Try to read image dimensions
                    try:
                        image_data = docx_zip.read(file_name)
                        if file_name.lower().endswith('.png'):
                            # PNG header: width and height at bytes 16-23
                            if len(image_data) > 24:
                                width = struct.unpack('>I', image_data[16:20])[0]
                                height = struct.unpack('>I', image_data[20:24])[0]
                                max_width = max(max_width, width)
                                max_height = max(max_height, height)
                        elif file_name.lower().endswith(('.jpg', '.jpeg')):
                            width, height = parse_jpeg_dimensions(image_data)
                            max_width = max(max_width, width)
                            max_height = max(max_height, height)
                    except Exception:
                        pass

        # Determine recommended DPI based on max image dimensions
        recommended_dpi = 300
        if max_width > 3000 or max_height > 3000:
            recommended_dpi = 600
        elif max_width > 1500 or max_height > 1500:
            recommended_dpi = 300
        else:
            recommended_dpi = 150

        return {
            "imageCount": image_count,
            "maxWidth": max_width,
            "maxHeight": max_height,
            "totalImageSize": total_image_size,
            "totalImageSizeFormatted": format_size(total_image_size),
            "recommendedDpi": recommended_dpi,
            "hasHighResImages": max_width > 1500 or max_height > 1500,
        }
    except Exception as e:
        logger.warning(f"Failed to analyze DOCX images: {e}")
        return {
            "imageCount": 0,
            "maxWidth": 0,
            "maxHeight": 0,
            "totalImageSize": 0,
            "totalImageSizeFormatted": "0 B",
            "recommendedDpi": 300,
            "hasHighResImages": False,
        }


def _apply_page_setup(doc, options: dict):
    """Apply page setup options to the Word document before export."""
    try:
        setup = doc.Sections(1).PageSetup

        page_size = options.get("pageSize", "A4")
        if page_size == "Letter":
            setup.PageWidth = 21.59 * 28.35
            setup.PageHeight = 27.94 * 28.35
        elif page_size == "Legal":
            setup.PageWidth = 21.59 * 28.35
            setup.PageHeight = 35.56 * 28.35
        else:  # A4
            setup.PageWidth = 21.0 * 28.35
            setup.PageHeight = 29.7 * 28.35

        orientation = options.get("orientation", "Portrait")
        if orientation in ("Landscape", "横向"):
            setup.Orientation = 1  # wdOrientLandscape
        else:
            setup.Orientation = 0  # wdOrientPortrait
    except Exception:
        pass  # Non-critical — proceed with defaults


def _apply_font_settings(word, doc, options: dict):
    """Apply font embedding settings to the Word document.

    Controls whether fonts are embedded in the exported PDF.
    """
    embed_fonts = options.get("embedFonts", True)
    
    try:
        # Set font embedding option
        word.Options.EmbedTrueTypeFonts = embed_fonts
        
        # Also set subset embedding (only embed used characters)
        if embed_fonts:
            word.Options.EmbedSubsetFonts = True
    except Exception:
        pass
    
    # Apply to document's font settings if possible
    try:
        if not embed_fonts:
            # If not embedding fonts, we can't really prevent it at document level
            # Word will use system fonts, but the PDF may still contain font references
            pass
    except Exception:
        pass


def _apply_image_quality_word(word, doc, options: dict):
    """Configure Word image compression settings before PDF export.

    Controls picture resolution and compression to match user requirements.
    Supports lossless mode for maximum quality preservation.
    """
    image_dpi = parse_image_dpi(options)
    lossless = options.get("losslessImages", False) if options else False

    # Disable global compression
    try:
        word.Options.CompressPictures = False
    except Exception:
        pass

    # Process each inline picture
    try:
        for section in doc.Sections:
            for shape in section.Range.InlineShapes:
                if shape.Type == 1:  # wdInlineShapePicture
                    # Reset to original quality
                    shape.Range.InlineShapes(1).ResetPicture()

                    if lossless:
                        # Lossless: set print quality and maximum DPI
                        # 0 = wdPictureCompressionPrint
                        # 9999 = no scaling (preserve original size)
                        try:
                            shape.Range.InlineShapes(1).ScaleWidth = 100
                            shape.Range.InlineShapes(1).ScaleHeight = 100
                        except Exception:
                            pass

        # Also process floating shapes (not inline)
        for section in doc.Sections:
            for shape in section.Range.ShapeRange:
                if shape.Type == 6:  # msoPicture
                    try:
                        if lossless:
                            shape.ScaleWidth = 100
                            shape.ScaleHeight = 100
                    except Exception:
                        pass
    except Exception:
        pass

    # Set image resolution for PDF export
    try:
        # For Word 2010+, we can set the PIF resolution
        # This controls the maximum DPI for embedded images
        if image_dpi >= 600:
            # For 600 DPI, ensure no downsampling
            word.Options.EmbedTrueTypeFonts = True
            word.Options.AutosaveInterval = 0  # Disable autosave to prevent compression
    except Exception:
        pass


def get_pdf_metadata(pdf_path: str) -> dict:
    """Extract metadata from a PDF file using PyMuPDF."""
    try:
        import pymupdf
        doc = pymupdf.open(pdf_path)
        pages = len(doc)
        meta = doc.metadata or {}
        doc.close()
    except Exception:
        pages = 0
        meta = {}

    file_size = os.path.getsize(pdf_path) if os.path.isfile(pdf_path) else 0

    return {
        "pages": pages,
        "size": file_size,
        "sizeFormatted": format_size(file_size),
        "pdfVersion": meta.get("format", "PDF 1.7"),
        "title": meta.get("title", ""),
        "author": meta.get("author", ""),
    }


def format_size(bytes_val: int) -> str:
    if bytes_val < 1024:
        return f"{bytes_val} B"
    if bytes_val < 1024 * 1024:
        return f"{bytes_val / 1024:.1f} KB"
    return f"{bytes_val / (1024 * 1024):.1f} MB"


def parse_image_dpi(options: dict | None, default=300) -> int:
    """Extract imageDpi from options dict, returning default on parse failure."""
    if options:
        try:
            return int(options.get("imageDpi", default))
        except (ValueError, TypeError):
            pass
    return default


def delete_job_files(job_id: str, row):
    """Remove all on-disk files associated with a job."""
    for path in [row["input_path"], row["output_path"]]:
        if path and os.path.isfile(path):
            try:
                os.remove(path)
            except OSError:
                pass
    # Clean up _media directory for markdown jobs
    if row["conversion_type"] == "to_markdown" and row["output_path"]:
        base_name = row["name"].rsplit(".", 1)[0]
        media_dir = os.path.join(os.path.dirname(row["output_path"]), base_name + "_media")
        if os.path.isdir(media_dir):
            try:
                shutil.rmtree(media_dir)
            except OSError:
                pass
    for suffix in ("", "_source"):
        temp_path = os.path.join(TEMP_DIR, f"{job_id}{suffix}.pdf")
        if os.path.isfile(temp_path):
            try:
                os.remove(temp_path)
            except OSError:
                pass


def _estimate_progress(job_id: str, input_size: int):
    """Estimate conversion progress based on file size and elapsed time."""
    # Estimate expected duration based on file size
    if input_size < 100 * 1024:
        expected_s = 3
    elif input_size < 1024 * 1024:
        expected_s = 6
    elif input_size < 10 * 1024 * 1024:
        expected_s = 12
    else:
        expected_s = 20

    start = time.monotonic()
    while True:
        time.sleep(0.4)
        with get_db() as db:
            row = db.execute("SELECT status FROM jobs WHERE id=?", (job_id,)).fetchone()
            if not row or row["status"] != "processing":
                return
            elapsed = time.monotonic() - start
            ratio = min(elapsed / expected_s, 1.0)
            progress = int(95 * (1 - (1 - ratio) ** 2))
            db.execute("UPDATE jobs SET progress=? WHERE id=?", (progress, job_id))
            db.commit()


def _run_conversion(job_id: str, input_path: str, output_dir: str, options: dict | None):
    """Background thread: run conversion and update job status."""
    # Read conversion type from DB first to determine if COM is needed
    with get_db() as db:
        row = db.execute("SELECT conversion_type FROM jobs WHERE id=?", (job_id,)).fetchone()
        conv_type = row["conversion_type"] if row else "doc_to_pdf"

    # Only initialize COM for non-PDF inputs that need Word
    use_com = not input_path.lower().endswith('.pdf')
    if use_com:
        import pythoncom
        pythoncom.CoInitialize()

    try:
        input_size = os.path.getsize(input_path)

        # Use timestamp-based output directory
        with _output_lock:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_dir = os.path.join(OUTPUT_DIR, timestamp)
            os.makedirs(output_dir, exist_ok=True)

        # Start progress estimation
        progress_thread = threading.Thread(
            target=_estimate_progress, args=(job_id, input_size), daemon=True
        )
        progress_thread.start()

        if conv_type == "pdf_to_docx":
            # PDF → DOCX conversion
            docx_path = convert_pdf_to_docx(input_path, output_dir, options)

            # Wait for progress estimation to finish
            progress_thread.join(timeout=2)

            meta = get_docx_metadata(docx_path)

            # Copy source PDF to temp for preview access
            try:
                shutil.copy2(input_path, os.path.join(TEMP_DIR, f"{job_id}_source.pdf"))
            except Exception:
                pass

            with get_db() as db:
                db.execute(
                    """UPDATE jobs SET
                        status='done', progress=100, output_path=?,
                        docx_size=?, docx_size_formatted=?,
                        finished_at=?
                        WHERE id=?""",
                    (docx_path, meta["size"], meta["sizeFormatted"],
                     datetime.now().isoformat(), job_id),
                )
                db.commit()
        elif conv_type == "to_markdown":
            # PDF/DOC/DOCX → Markdown conversion
            intermediate_pdf = None
            if input_path.lower().endswith('.pdf'):
                # PDF: direct conversion
                md_path, media_dir = convert_to_markdown(input_path, output_dir, options)
            else:
                # DOC/DOCX: first convert to PDF, then to Markdown
                temp_pdf = convert_doc_to_pdf(input_path, output_dir, options)
                intermediate_pdf = temp_pdf
                try:
                    md_path, media_dir = convert_to_markdown(temp_pdf, output_dir, options)
                finally:
                    pass  # Keep temp PDF for preview, cleanup below

            # Wait for progress estimation to finish
            progress_thread.join(timeout=2)

            md_size = os.path.getsize(md_path) if os.path.isfile(md_path) else 0

            # Copy source to temp for preview access
            try:
                source_pdf = os.path.join(TEMP_DIR, f"{job_id}_source.pdf")
                if input_path.lower().endswith('.pdf'):
                    shutil.copy2(input_path, source_pdf)
                elif intermediate_pdf and os.path.isfile(intermediate_pdf):
                    shutil.copy2(intermediate_pdf, source_pdf)
            except Exception:
                pass

            # Cleanup intermediate PDF (after preview copy)
            if intermediate_pdf:
                try:
                    os.remove(intermediate_pdf)
                except OSError:
                    pass

            with get_db() as db:
                db.execute(
                    """UPDATE jobs SET
                        status='done', progress=100, output_path=?,
                        markdown_size=?, markdown_size_formatted=?,
                        finished_at=?
                        WHERE id=?""",
                    (md_path, md_size, format_size(md_size),
                     datetime.now().isoformat(), job_id),
                )
                db.commit()
        else:
            # DOC/DOCX → PDF conversion
            pdf_path = convert_doc_to_pdf(input_path, output_dir, options)

            # Wait for progress estimation to finish
            progress_thread.join(timeout=2)

            meta = get_pdf_metadata(pdf_path)

            # Copy to temp for preview access
            try:
                shutil.copy2(pdf_path, os.path.join(TEMP_DIR, f"{job_id}.pdf"))
            except Exception:
                pass

            with get_db() as db:
                db.execute(
                    """UPDATE jobs SET
                        status='done', progress=100, output_path=?,
                        pages=?, pdf_size=?, pdf_size_formatted=?, pdf_version=?,
                        finished_at=?
                        WHERE id=?""",
                    (pdf_path, meta["pages"], meta["size"],
                     meta["sizeFormatted"], meta["pdfVersion"],
                     datetime.now().isoformat(), job_id),
                )
                db.commit()
    except Exception as e:
        with get_db() as db:
            db.execute(
                "UPDATE jobs SET status='error', error=? WHERE id=?",
                (str(e), job_id),
            )
            db.commit()
    finally:
        if use_com:
            pythoncom.CoUninitialize()


# ── Flask routes ──


@app.route("/")
def index():
    html_path = os.path.join(BUNDLE_DIR, "index.html")
    with open(html_path, "r", encoding="utf-8") as f:
        html = f.read()

    # Inject custom font @font-face + CSS variable so it applies on first paint
    font_css = ""
    meta_path = os.path.join(FONT_DIR, "font.json")
    if os.path.isfile(meta_path):
        try:
            with open(meta_path, "r", encoding="utf-8") as mf:
                meta = json.load(mf)
            if meta.get("name"):
                name = meta["name"]
                ext = os.path.splitext(meta["filename"])[1].lower()
                fmt_map = {".ttf": "truetype", ".otf": "opentype", ".woff": "woff", ".woff2": "woff2"}
                font_fmt = fmt_map.get(ext, "truetype")
                font_css = (
                    f'<link rel="preload" href="/api/font-file" as="font" crossorigin>'
                    f'<style>'
                    f'@font-face{{font-family:\'{name}\';src:url("/api/font-file") format("{font_fmt}");font-display:swap}}'
                    f':root{{--font-custom:\'{name}\'}}'
                    f'</style>'
                )
        except Exception:
            pass

    html = html.replace("</head>", f"{font_css}</head>", 1)
    return html


@app.route("/favicon.ico")
def serve_favicon():
    return send_from_directory(BUNDLE_DIR, "icon.png", mimetype="image/png")

@app.route("/icon.png")
def serve_icon():
    return send_from_directory(BUNDLE_DIR, "icon.png", mimetype="image/png")


@app.route("/icon01.png")
def serve_icon01():
    return send_from_directory(BUNDLE_DIR, "icon01.png", mimetype="image/png")


@app.route("/pdf.min.js")
def serve_pdf_js():
    return send_from_directory(BUNDLE_DIR, "pdf.min.js", mimetype="application/javascript")


@app.route("/pdf.worker.min.js")
def serve_pdf_worker():
    return send_from_directory(BUNDLE_DIR, "pdf.worker.min.js", mimetype="application/javascript")


@app.route("/styles.css")
def serve_styles():
    return send_from_directory(BUNDLE_DIR, "styles.css", mimetype="text/css")


@app.route("/app.js")
def serve_app_js():
    return send_from_directory(BUNDLE_DIR, "app.js", mimetype="application/javascript")

@app.route("/HELP.mp4")
def serve_help_video():
    return send_from_directory(BASE_DIR, "HELP.mp4", mimetype="video/mp4")


@app.route("/READ.mp3")
def serve_read_audio():
    return send_from_directory(BASE_DIR, "READ.mp3", mimetype="audio/mpeg")


@app.route("/help_docs/<path:filename>")
def serve_help_docs(filename):
    """Serve help documentation files."""
    help_dir = os.path.join(BASE_DIR, "help_docs")
    return send_from_directory(help_dir, filename)


@app.route("/api/engine")
def engine_status():
    ocr_status = _check_ocr_engines()
    return jsonify({
        "wordAvailable": is_word_available(),
        "ocr": ocr_status,
    })


@app.route("/api/upload", methods=["POST"])
def upload():
    if "files" not in request.files:
        return jsonify({"error": "未选择文件"}), 400

    files = request.files.getlist("files")
    if not files:
        return jsonify({"error": "未选择文件"}), 400

    jobs = []
    mode = request.form.get('mode', '')
    with get_db() as db:
        for f in files:
            name_lower = f.filename.lower()
            if mode == 'to_markdown':
                if name_lower.endswith(('.pdf', '.doc', '.docx')):
                    conv_type = "to_markdown"
                else:
                    continue
            elif name_lower.endswith(".pdf"):
                conv_type = "pdf_to_docx"
            elif name_lower.endswith(".doc") or name_lower.endswith(".docx"):
                conv_type = "doc_to_pdf"
            else:
                continue

            job_id = str(uuid.uuid4())
            safe_name = f.filename
            input_path = os.path.join(UPLOAD_DIR, f"{job_id}_{safe_name}")
            f.save(input_path)
            file_size = os.path.getsize(input_path)

            # Analyze DOCX images for quality recommendations
            image_analysis = None
            if conv_type == "doc_to_pdf":
                image_analysis = analyze_docx_images(input_path)

            db.execute(
                """INSERT INTO jobs (id, name, conversion_type, input_path, input_size, input_size_formatted, created_at)
                   VALUES (?, ?, ?, ?, ?, ?, ?)""",
                (job_id, safe_name, conv_type, input_path, file_size,
                 format_size(file_size), datetime.now().isoformat()),
            )

            jobs.append({
                "id": job_id,
                "name": safe_name,
                "size": file_size,
                "sizeFormatted": format_size(file_size),
                "conversionType": conv_type,
                "imageAnalysis": image_analysis,
            })
        db.commit()

    if not jobs:
        return jsonify({"error": "未找到有效的文件（支持 DOC/DOCX/PDF）"}), 400

    return jsonify({"jobs": jobs})


@app.route("/api/convert/<job_id>", methods=["POST"])
def convert(job_id):
    with get_db() as db:
        row = db.execute("SELECT * FROM jobs WHERE id=?", (job_id,)).fetchone()
        if not row:
            return jsonify({"error": "任务不存在"}), 404

        if row["status"] not in ("pending", "error"):
            return jsonify({"error": "任务已在处理中或已完成"}), 400

        if row["conversion_type"] != "pdf_to_docx" and not row["input_path"].lower().endswith('.pdf') and not is_word_available():
            db.execute(
                "UPDATE jobs SET status='error', error=? WHERE id=?",
                ("Microsoft Word 不可用", job_id),
            )
            db.commit()
            return jsonify({"error": "Microsoft Word 不可用"}), 500

        db.execute(
            "UPDATE jobs SET status='processing', progress=0, error=NULL WHERE id=?",
            (job_id,),
        )
        db.commit()
        input_path = row["input_path"]

    options = None
    try:
        options = request.get_json(silent=True)
    except Exception:
        pass

    thread = threading.Thread(
        target=_run_conversion,
        args=(job_id, input_path, OUTPUT_DIR, options),
        daemon=True,
    )
    thread.start()

    return jsonify({"status": "processing", "jobId": job_id})


@app.route("/api/reconvert/<job_id>", methods=["POST"])
def reconvert(job_id):
    with get_db() as db:
        row = db.execute("SELECT * FROM jobs WHERE id=?", (job_id,)).fetchone()
        if not row:
            return jsonify({"error": "任务不存在"}), 404

        if row["status"] == "processing":
            return jsonify({"error": "任务正在处理中"}), 400

        if row["conversion_type"] != "pdf_to_docx" and not row["input_path"].lower().endswith('.pdf') and not is_word_available():
            return jsonify({"error": "Microsoft Word 不可用"}), 500

        if not row["input_path"] or not os.path.isfile(row["input_path"]):
            return jsonify({"error": "原始文件不存在"}), 404

        db.execute(
            "UPDATE jobs SET status='processing', progress=0, error=NULL, output_path=NULL WHERE id=?",
            (job_id,),
        )
        db.commit()

        # Invalidate stale DOCX preview cache from previous conversion
        preview_pdf = os.path.join(TEMP_DIR, f"{job_id}_preview.pdf")
        try:
            os.remove(preview_pdf)
        except OSError:
            pass

        input_path = row["input_path"]

    options = None
    try:
        options = request.get_json(silent=True)
    except Exception:
        pass

    thread = threading.Thread(
        target=_run_conversion,
        args=(job_id, input_path, OUTPUT_DIR, options),
        daemon=True,
    )
    thread.start()

    return jsonify({"status": "processing", "jobId": job_id})


@app.route("/api/status/<job_id>")
def status(job_id):
    with get_db() as db:
        row = db.execute("SELECT * FROM jobs WHERE id=?", (job_id,)).fetchone()
    if not row:
        return jsonify({"error": "任务不存在"}), 404

    result = {
        "status": row["status"],
        "progress": row["progress"],
        "conversionType": row["conversion_type"],
        "error": row["error"],
    }
    if row["conversion_type"] == "pdf_to_docx":
        result["docxSize"] = row["docx_size"]
        result["docxSizeFormatted"] = row["docx_size_formatted"]
    elif row["conversion_type"] == "to_markdown":
        result["markdownSize"] = row["markdown_size"]
        result["markdownSizeFormatted"] = row["markdown_size_formatted"]
    else:
        result["pages"] = row["pages"]
        result["pdfSize"] = row["pdf_size"]
    return jsonify(result)


@app.route("/api/download/<job_id>")
def download(job_id):
    with get_db() as db:
        row = db.execute("SELECT * FROM jobs WHERE id=?", (job_id,)).fetchone()
    if not row or row["status"] != "done":
        return jsonify({"error": "文件未就绪"}), 404

    output_path = row["output_path"]
    if not output_path or not os.path.isfile(output_path):
        return jsonify({"error": "文件不存在"}), 404

    base_name = row["name"].rsplit(".", 1)[0]
    if row["conversion_type"] == "pdf_to_docx":
        download_name = base_name + ".docx"
        mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif row["conversion_type"] == "to_markdown":
        # Zip the .md file and _media folder together
        # Derive base_name from actual output file path (has UUID prefix), not row["name"]
        md_path = output_path
        md_basename = os.path.splitext(os.path.basename(md_path))[0]
        media_dir = os.path.join(os.path.dirname(md_path), md_basename + "_media")
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.write(md_path, base_name + ".md")
            if os.path.isdir(media_dir):
                for root, dirs, files in os.walk(media_dir):
                    for fn in files:
                        fp = os.path.join(root, fn)
                        arcname = os.path.join(base_name + "_media", os.path.relpath(fp, media_dir))
                        zf.write(fp, arcname)
        buf.seek(0)
        return send_file(
            buf,
            as_attachment=True,
            download_name=base_name + ".zip",
            mimetype="application/zip",
        )
    else:
        download_name = base_name + ".pdf"
        mimetype = "application/pdf"

    return send_file(
        output_path,
        as_attachment=True,
        download_name=download_name,
        mimetype=mimetype,
    )


@app.route("/api/preview/<job_id>")
def preview(job_id):
    with get_db() as db:
        row = db.execute("SELECT * FROM jobs WHERE id=?", (job_id,)).fetchone()
    if not row or row["status"] != "done":
        return jsonify({"error": "文件未就绪"}), 404

    temp_path = os.path.join(TEMP_DIR, f"{job_id}.pdf")
    source_path = os.path.join(TEMP_DIR, f"{job_id}_source.pdf")
    output_path = row["output_path"]

    for path in [temp_path, source_path, output_path]:
        if path and os.path.isfile(path):
            try:
                return send_file(path, mimetype="application/pdf", as_attachment=False)
            except (OSError, PermissionError):
                continue

    return jsonify({"error": "文件不存在"}), 404


@app.route("/api/markdown/<job_id>")
def preview_markdown(job_id):
    """Serve markdown file content for preview rendering."""
    with get_db() as db:
        row = db.execute("SELECT output_path FROM jobs WHERE id=?", (job_id,)).fetchone()
    if not row or not row["output_path"] or not os.path.isfile(row["output_path"]):
        return jsonify({"error": "文件不存在"}), 404
    try:
        with open(row["output_path"], "r", encoding="utf-8", errors="replace") as f:
            content = f.read()
        return jsonify({"content": content})
    except (OSError, PermissionError):
        return jsonify({"error": "读取失败"}), 500


@app.route("/api/image/<job_id>/<path:filename>")
def serve_markdown_image(job_id, filename):
    """Serve image files from markdown _media directory for preview."""
    with get_db() as db:
        row = db.execute("SELECT output_path FROM jobs WHERE id=?", (job_id,)).fetchone()
    if not row or not row["output_path"]:
        return jsonify({"error": "任务不存在"}), 404
    output_dir = os.path.dirname(row["output_path"])
    # Images are stored in {md_basename}_media/ subdirectory
    md_basename = os.path.splitext(os.path.basename(row["output_path"]))[0]
    media_dir = os.path.join(output_dir, md_basename + "_media")
    # Look in media_dir first, then output_dir as fallback
    for search_dir in [media_dir, output_dir]:
        image_path = os.path.join(search_dir, filename)
        if os.path.isfile(image_path) and os.path.realpath(image_path).startswith(os.path.realpath(output_dir)):
            ext = os.path.splitext(filename)[1].lower()
            mimetypes = {'.png': 'image/png', '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg',
                         '.gif': 'image/gif', '.webp': 'image/webp', '.svg': 'image/svg+xml'}
            return send_file(image_path, mimetype=mimetypes.get(ext, 'application/octet-stream'))
    return jsonify({"error": "文件不存在"}), 404


@app.route("/api/preview-source/<job_id>")
def preview_source(job_id):
    """Serve the source PDF for preview (used by PDF→DOCX jobs)."""
    with get_db() as db:
        row = db.execute(
            "SELECT input_path, conversion_type FROM jobs WHERE id=?", (job_id,)
        ).fetchone()
    if not row or row["conversion_type"] not in ("pdf_to_docx", "to_markdown"):
        return jsonify({"error": "任务不存在"}), 404

    temp_path = os.path.join(TEMP_DIR, f"{job_id}_source.pdf")
    input_path = row["input_path"]

    for path in [temp_path, input_path]:
        if path and os.path.isfile(path):
            try:
                return send_file(path, mimetype="application/pdf", as_attachment=False)
            except (OSError, PermissionError):
                continue

    return jsonify({"error": "源文件不存在"}), 404


@app.route("/api/preview-docx/<job_id>")
def preview_docx(job_id):
    """Convert DOCX output to PDF for preview."""
    with get_db() as db:
        row = db.execute(
            "SELECT output_path, conversion_type, status FROM jobs WHERE id=?", (job_id,)
        ).fetchone()

    if not row:
        return jsonify({"error": "任务不存在"}), 404

    if row["conversion_type"] != "pdf_to_docx":
        return jsonify({"error": "仅支持PDF转DOCX任务的预览"}), 400

    if row["status"] != "done":
        return jsonify({"error": "转换尚未完成"}), 400

    docx_path = row["output_path"]
    if not docx_path or not os.path.isfile(docx_path):
        return jsonify({"error": "DOCX文件不存在"}), 404

    # Check if preview PDF already exists
    preview_pdf_path = os.path.join(TEMP_DIR, f"{job_id}_preview.pdf")
    if os.path.isfile(preview_pdf_path):
        try:
            return send_file(preview_pdf_path, mimetype="application/pdf", as_attachment=False)
        except (OSError, PermissionError):
            pass  # Regenerate if send fails

    # Convert DOCX to PDF for preview using Word COM
    if not is_word_available():
        return jsonify({"error": "Microsoft Word 不可用，无法生成预览"}), 500

    word = None
    doc = None
    try:
        import pythoncom
        import win32com.client
        pythoncom.CoInitialize()

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0

        doc = word.Documents.Open(os.path.abspath(docx_path))

        # Configure image quality for preview (use maximum quality)
        _apply_image_quality_word(word, doc, {"imageDpi": 600, "losslessImages": True})

        # Export as PDF with high quality image settings
        # Use ExportAsFixedFormat for better control over PDF quality
        doc.ExportAsFixedFormat(
            OutputFileName=os.path.abspath(preview_pdf_path),
            ExportFormat=17,  # 17 = PDF
            OpenAfterExport=False,
            OptimizeFor=0,  # 0 = Standard (higher quality), 1 = Minimum size
            Range=0,  # 0 = All pages
            From=1,
            To=0,
            Item=0,  # 0 = Document content
            IncludeDocProps=True,
            KeepIRM=True,
            CreateBookmarks=1,  # 1 = Word bookmarks
            DocStructureTags=True,
            BitmapMissingFonts=True,
            UseISO19005_1=False
        )

        doc.Close()
        doc = None

        word.Quit()
        word = None

        pythoncom.CoUninitialize()

        if os.path.isfile(preview_pdf_path):
            return send_file(preview_pdf_path, mimetype="application/pdf", as_attachment=False)
        else:
            return jsonify({"error": "预览生成失败"}), 500

    except Exception as e:
        logger.error(f"DOCX preview generation failed: {e}")
        if doc:
            try:
                doc.Close(SaveChanges=False)
            except Exception:
                pass
        if word:
            try:
                word.Quit()
            except Exception:
                pass
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass
        return jsonify({"error": f"预览生成失败: {str(e)}"}), 500


@app.route("/api/download-all")
def download_all():
    conv_type = request.args.get("type")
    with get_db() as db:
        if conv_type in ("doc_to_pdf", "pdf_to_docx", "to_markdown"):
            rows = db.execute(
                "SELECT * FROM jobs WHERE status='done' AND output_path IS NOT NULL AND conversion_type=?",
                (conv_type,),
            ).fetchall()
        else:
            rows = db.execute(
                "SELECT * FROM jobs WHERE status='done' AND output_path IS NOT NULL"
            ).fetchall()

    if not rows:
        return jsonify({"error": "没有可下载的文件"}), 404

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        seen = set()
        for row in rows:
            if not row["output_path"] or not os.path.isfile(row["output_path"]):
                continue
            base_name = row["name"].rsplit(".", 1)[0]
            if row["conversion_type"] == "to_markdown":
                # Zip .md + _media folder
                # Derive md_basename from actual output file path (has UUID prefix)
                md_basename = os.path.splitext(os.path.basename(row["output_path"]))[0]
                arcname = base_name + ".md"
                if arcname in seen:
                    base, ext = arcname.rsplit(".", 1)
                    i = 2
                    while f"{base}_{i}.{ext}" in seen:
                        i += 1
                    arcname = f"{base}_{i}.{ext}"
                seen.add(arcname)
                zf.write(row["output_path"], arcname)
                media_dir = os.path.join(os.path.dirname(row["output_path"]), md_basename + "_media")
                if os.path.isdir(media_dir):
                    for root, dirs, files in os.walk(media_dir):
                        for fn in files:
                            fp = os.path.join(root, fn)
                            media_arcname = os.path.join(base_name + "_media", os.path.relpath(fp, media_dir))
                            zf.write(fp, media_arcname)
                continue
            arcname = base_name + (".docx" if row["conversion_type"] == "pdf_to_docx" else ".pdf")
            if arcname in seen:
                base, ext = arcname.rsplit(".", 1)
                i = 2
                while f"{base}_{i}.{ext}" in seen:
                    i += 1
                arcname = f"{base}_{i}.{ext}"
            seen.add(arcname)
            zf.write(row["output_path"], arcname)

    buf.seek(0)
    return send_file(
        buf,
        as_attachment=True,
        download_name="docflow_batch.zip",
        mimetype="application/zip",
    )


@app.route("/api/history")
def history():
    conv_type = request.args.get("type")
    with get_db() as db:
        if conv_type in ("doc_to_pdf", "pdf_to_docx", "to_markdown"):
            rows = db.execute(
                "SELECT * FROM jobs WHERE conversion_type=? AND status IN ('done','error') ORDER BY created_at DESC",
                (conv_type,),
            ).fetchall()
        else:
            rows = db.execute(
                "SELECT * FROM jobs WHERE status IN ('done','error') ORDER BY created_at DESC"
            ).fetchall()

    items = []
    for r in rows:
        item = {
            "id": r["id"],
            "name": r["name"],
            "status": r["status"],
            "progress": r["progress"],
            "conversionType": r["conversion_type"],
            "inputSize": r["input_size"],
            "inputSizeFormatted": r["input_size_formatted"],
            "error": r["error"],
            "finishedAt": r["finished_at"],
        }
        if r["conversion_type"] == "pdf_to_docx":
            item["docxSize"] = r["docx_size"]
            item["docxSizeFormatted"] = r["docx_size_formatted"]
        elif r["conversion_type"] == "to_markdown":
            item["markdownSize"] = r["markdown_size"]
            item["markdownSizeFormatted"] = r["markdown_size_formatted"]
        else:
            item["pdfSize"] = r["pdf_size"]
            item["pdfSizeFormatted"] = r["pdf_size_formatted"]
            item["pages"] = r["pages"]
        items.append(item)

    return jsonify({"history": items})


@app.route("/api/active-jobs")
def active_jobs():
    conv_type = request.args.get("type")
    with get_db() as db:
        if conv_type in ("doc_to_pdf", "pdf_to_docx", "to_markdown"):
            rows = db.execute(
                "SELECT id, name, status, progress, conversion_type, input_size, input_size_formatted FROM jobs WHERE conversion_type=? AND status IN ('pending','processing')",
                (conv_type,),
            ).fetchall()
        else:
            rows = db.execute(
                "SELECT id, name, status, progress, conversion_type, input_size, input_size_formatted FROM jobs WHERE status IN ('pending','processing')"
            ).fetchall()
    return jsonify({"jobs": [{"id": r["id"], "name": r["name"], "status": r["status"], "progress": r["progress"], "conversionType": r["conversion_type"], "size": r["input_size"], "sizeFormatted": r["input_size_formatted"]} for r in rows]})


@app.route("/api/history/<job_id>", methods=["DELETE"])
def delete_history(job_id):
    with get_db() as db:
        row = db.execute("SELECT * FROM jobs WHERE id=?", (job_id,)).fetchone()
        if not row:
            return jsonify({"error": "记录不存在"}), 404
        delete_job_files(job_id, row)
        db.execute("DELETE FROM jobs WHERE id=?", (job_id,))
        db.commit()
    return jsonify({"ok": True})


@app.route("/api/history", methods=["DELETE"])
def clear_history():
    conv_type = request.args.get("type")
    with get_db() as db:
        if conv_type in ("doc_to_pdf", "pdf_to_docx", "to_markdown"):
            rows = db.execute("SELECT * FROM jobs WHERE conversion_type=?", (conv_type,)).fetchall()
        else:
            rows = db.execute("SELECT * FROM jobs").fetchall()

        for row in rows:
            delete_job_files(row["id"], row)

        if conv_type in ("doc_to_pdf", "pdf_to_docx", "to_markdown"):
            db.execute("DELETE FROM jobs WHERE conversion_type=?", (conv_type,))
        else:
            db.execute("DELETE FROM jobs")
        db.commit()
    return jsonify({"ok": True})


@app.route("/api/metadata/<job_id>")
def metadata(job_id):
    with get_db() as db:
        row = db.execute("SELECT * FROM jobs WHERE id=?", (job_id,)).fetchone()

    if not row or row["status"] != "done":
        return jsonify({"error": "记录不存在或未完成"}), 404

    if row["conversion_type"] == "pdf_to_docx":
        output_path = row["output_path"]
        if output_path and os.path.isfile(output_path):
            return jsonify(get_docx_metadata(output_path))
        return jsonify({
            "size": row["docx_size"],
            "sizeFormatted": row["docx_size_formatted"],
        })
    elif row["conversion_type"] == "to_markdown":
        result = {
            "size": row["markdown_size"],
            "sizeFormatted": row["markdown_size_formatted"],
        }
        # Try to get page count from source PDF
        source_path = os.path.join(TEMP_DIR, f"{job_id}_source.pdf")
        if not os.path.isfile(source_path):
            source_path = row["input_path"]
        if source_path and os.path.isfile(source_path) and source_path.lower().endswith('.pdf'):
            try:
                doc = fitz.open(source_path)
                result["pages"] = len(doc)
                doc.close()
            except Exception:
                pass
        return jsonify(result)
    else:
        output_path = row["output_path"]
        if output_path and os.path.isfile(output_path):
            return jsonify(get_pdf_metadata(output_path))
        return jsonify({
            "pages": row["pages"],
            "size": row["pdf_size"],
            "sizeFormatted": row["pdf_size_formatted"],
            "pdfVersion": row["pdf_version"],
        })


# ── Mindmap Generation API ──

KMIND_SCRIPT = os.path.join(BASE_DIR, "kmind-markdown-to-mindmap-0.1.0", "scripts", "kmind-render.mjs")
# CJK-safe font stack — ensures consistent metrics for Chinese/Japanese/Korean text
_CJK_FONT_FAMILY = (
    '"Microsoft YaHei", "PingFang SC", "Hiragino Sans GB", '
    '"Noto Sans CJK SC", "Source Han Sans SC", '
    '"ui-sans-serif", "system-ui", "sans-serif"'
)


def _patch_svg_cjk_fonts(svg_path):
    """Replace kmind's default font-family with CJK-compatible stack and fix layout."""
    try:
        with open(svg_path, "r", encoding="utf-8") as f:
            content = f.read()
        # Replace font-family attribute values
        original = "ui-sans-serif, system-ui, sans-serif"
        content = content.replace(f'font-family="{original}"', f'font-family="{_CJK_FONT_FAMILY}"')
        # Also handle CSS-style fontFamily declarations (in foreignObject style blocks)
        content = content.replace(
            f'fontFamily: "{original}"',
            f'fontFamily: "{_CJK_FONT_FAMILY}"',
        )
        content = content.replace(
            f"fontFamily: '{original}'",
            f"fontFamily: '{_CJK_FONT_FAMILY}'",
        )
        # Inject CSS to fix title/body overlap and improve space utilization
        layout_fix = """
  .km-node-body { overflow: hidden; }
  .km-node-region-right { overflow: hidden; min-width: 0; }
  .km-node-notes { overflow: hidden; word-break: break-word; overflow-wrap: anywhere; }
  .km-node-region-main { min-width: 0; overflow: hidden; }
  .km-node-richtext { overflow: hidden; }
"""
        content = content.replace("</style>", layout_fix + "</style>")
        with open(svg_path, "w", encoding="utf-8") as f:
            f.write(content)
    except Exception:
        pass  # best-effort; don't fail the whole request


def _preprocess_md_for_mindmap(md_content):
    """Clean up markdown before mindmap generation to reduce node bloat."""
    import re
    text = md_content
    # Remove markdown image references — images can't render in mindmap nodes
    text = re.sub(r'!\[[^\]]*\]\([^)]*\)', '', text)
    # Remove HTML comment page markers
    text = re.sub(r'<!--\s*Page\s+\d+\s*-->', '', text)
    # Remove horizontal rules
    text = re.sub(r'^-{3,}\s*$', '', text, flags=re.MULTILINE)
    # Remove bullet symbols (⚫, •, ▪, etc.) — keep the text
    text = re.sub(r'^[\s]*[⚫•▪◦▸▹►◆◇○●■□△▽]\s*', '', text, flags=re.MULTILINE)
    # Collapse 3+ consecutive blank lines into 2
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Strip trailing whitespace per line
    text = re.sub(r'[ \t]+$', '', text, flags=re.MULTILINE)
    # Decode HTML entities that pymupdf sometimes emits
    text = text.replace('&lt;', '<').replace('&gt;', '>').replace('&amp;', '&')
    return text.strip()


@app.route("/api/mindmap-check")
def mindmap_check():
    """Pre-flight check: is Node.js installed and kmind script present?"""
    try:
        result = subprocess.run(
            ["node", "--version"], capture_output=True, timeout=5,
            creationflags=getattr(subprocess, "CREATE_NO_WINDOW", 0),
        )
        node_ok = result.returncode == 0
    except Exception:
        node_ok = False
    kmind_ok = os.path.isfile(KMIND_SCRIPT)
    return jsonify({"node": node_ok, "kmind": kmind_ok})


@app.route("/api/mindmap/<job_id>", methods=["POST"])
def generate_mindmap(job_id):
    """Generate a mind map (SVG/PNG) from a to_markdown job's output."""
    with get_db() as db:
        row = db.execute("SELECT * FROM jobs WHERE id=?", (job_id,)).fetchone()
    if not row or row["status"] != "done":
        return jsonify({"error": "文件未就绪"}), 404
    if row["conversion_type"] != "to_markdown":
        return jsonify({"error": "仅支持 Markdown 任务"}), 400

    md_path = row["output_path"]
    if not md_path or not os.path.isfile(md_path):
        return jsonify({"error": "Markdown 文件不存在"}), 404

    data = request.get_json(silent=True) or {}
    fmt = data.get("format", "svg")
    if fmt not in ("svg", "png"):
        fmt = "svg"
    png_scale = data.get("pngScale", 1)

    # Read markdown content
    try:
        with open(md_path, "r", encoding="utf-8") as f:
            md_content = f.read()
    except Exception as e:
        return jsonify({"error": f"读取 Markdown 失败: {e}"}), 500

    if not md_content.strip():
        return jsonify({"error": "Markdown 内容为空，无法生成脑图"}), 400

    # Preprocess: strip noise so kmind generates tighter nodes
    md_content = _preprocess_md_for_mindmap(md_content)

    # Check prerequisites
    if not os.path.isfile(KMIND_SCRIPT):
        return jsonify({"error": "脑图模块未找到"}), 500

    # Prepare temp output directory
    mindmap_dir = os.path.join(TEMP_DIR, f"mindmap_{job_id}")
    os.makedirs(mindmap_dir, exist_ok=True)
    out_ext = ".svg" if fmt == "svg" else ".png"
    out_path = os.path.join(mindmap_dir, f"mindmap{out_ext}")

    try:
        # Build command — larger viewport gives CJK text more room
        cmd = [
            "node", KMIND_SCRIPT,
            "render-markdown", "-",
            "--output", out_path,
            "--theme-preset", "kmind-material-3-slate",
            "--viewport-width", "2400",
            "--viewport-height", "1600",
        ]
        if fmt == "png":
            cmd.extend(["--png-scale", str(png_scale)])

        # Run kmind subprocess
        proc = subprocess.run(
            cmd,
            input=md_content.encode("utf-8"),
            capture_output=True,
            timeout=120,
            creationflags=getattr(subprocess, "CREATE_NO_WINDOW", 0),
        )

        # Check if output file was created
        if not os.path.isfile(out_path):
            stderr_msg = proc.stderr.decode("utf-8", errors="replace")[:500]
            return jsonify({"error": f"脑图生成失败: {stderr_msg or '未知错误'}"}), 500

        # Post-process SVG to fix CJK font overlap
        if fmt == "svg":
            _patch_svg_cjk_fonts(out_path)

        # Return the generated file
        base_name = row["name"].rsplit(".", 1)[0]
        download_name = f"{base_name}_mindmap{out_ext}"
        mimetype = "image/svg+xml" if fmt == "svg" else "image/png"

        # Read file into memory before cleanup
        with open(out_path, "rb") as f:
            file_data = f.read()

        return send_file(
            io.BytesIO(file_data),
            as_attachment=True,
            download_name=download_name,
            mimetype=mimetype,
        )
    except subprocess.TimeoutExpired:
        return jsonify({"error": "脑图生成超时（>120秒）"}), 500
    except Exception as e:
        return jsonify({"error": f"脑图生成失败: {e}"}), 500
    finally:
        # Cleanup temp files
        try:
            shutil.rmtree(mindmap_dir, ignore_errors=True)
        except Exception:
            pass


# ── Custom Font API ──

FONT_META_PATH = os.path.join(FONT_DIR, "font.json")


@app.route("/api/font", methods=["GET"])
def get_font():
    if os.path.isfile(FONT_META_PATH):
        with open(FONT_META_PATH, "r", encoding="utf-8") as f:
            meta = json.load(f)
        font_path = os.path.join(FONT_DIR, meta["filename"])
        if os.path.isfile(font_path):
            return jsonify({"name": meta["name"], "filename": meta["filename"]})
    return jsonify({"name": None})


@app.route("/api/font-file")
def serve_font_file():
    if not os.path.isfile(FONT_META_PATH):
        return jsonify({"error": "未设置自定义字体"}), 404
    with open(FONT_META_PATH, "r", encoding="utf-8") as f:
        meta = json.load(f)
    font_path = os.path.join(FONT_DIR, meta["filename"])
    if not os.path.isfile(font_path):
        return jsonify({"error": "字体文件不存在"}), 404
    ext = os.path.splitext(meta["filename"])[1].lower()
    mime_map = {".ttf": "font/ttf", ".otf": "font/otf", ".woff": "font/woff", ".woff2": "font/woff2"}
    resp = send_file(font_path, mimetype=mime_map.get(ext, "application/octet-stream"))
    resp.headers["Access-Control-Allow-Origin"] = "*"
    return resp


@app.route("/api/font", methods=["POST"])
def upload_font():
    if "font" not in request.files:
        return jsonify({"error": "未选择字体文件"}), 400
    f = request.files["font"]
    if not f.filename:
        return jsonify({"error": "未选择字体文件"}), 400

    ext = os.path.splitext(f.filename)[1].lower()
    if ext not in (".ttf", ".otf", ".woff", ".woff2"):
        return jsonify({"error": "不支持的字体格式"}), 400

    font_name = os.path.splitext(f.filename)[0]
    safe_filename = f"custom{ext}"

    # Remove old font files
    for old in os.listdir(FONT_DIR):
        old_path = os.path.join(FONT_DIR, old)
        if os.path.isfile(old_path):
            try:
                os.remove(old_path)
            except OSError:
                pass

    font_path = os.path.join(FONT_DIR, safe_filename)
    f.save(font_path)

    with open(FONT_META_PATH, "w", encoding="utf-8") as mf:
        json.dump({"name": font_name, "filename": safe_filename}, mf)

    return jsonify({"name": font_name, "filename": safe_filename})


@app.route("/api/font", methods=["DELETE"])
def delete_font():
    for fname in os.listdir(FONT_DIR):
        fpath = os.path.join(FONT_DIR, fname)
        if os.path.isfile(fpath):
            try:
                os.remove(fpath)
            except OSError:
                pass
    return jsonify({"ok": True})


# ── Music Player API ──

MUSIC_EXTS = ('.mp3', '.wav', '.ogg', '.flac', '.m4a')


@app.route("/api/music/list")
def list_music():
    tracks = []
    if os.path.isdir(MUSIC_DIR):
        for fname in sorted(os.listdir(MUSIC_DIR)):
            if fname.lower().endswith(MUSIC_EXTS):
                tracks.append({
                    "name": os.path.splitext(fname)[0],
                    "file": fname,
                })
    return jsonify({"tracks": tracks})


@app.route("/music/<path:filename>")
def serve_music(filename):
    return send_from_directory(MUSIC_DIR, filename, mimetype="audio/mpeg")


# ── Main ──

def _warmup_imports():
    """Pre-import heavy libraries so first conversion request isn't slow."""
    try:
        import pdf2docx  # noqa: F401
        import fitz  # noqa: F401
        import docx  # noqa: F401
    except ImportError:
        pass


if __name__ == "__main__":
    import webbrowser
    import threading
    port = 5000
    url = f"http://localhost:{port}"
    print("DocFlow 启动中...")
    _warmup_imports()
    print("依赖加载完成")
    print(f"打开浏览器访问: {url}")
    try:
        if getattr(sys, 'frozen', False):
            threading.Timer(1.5, lambda: webbrowser.open(url)).start()
        app.run(host="0.0.0.0", port=port, debug=not getattr(sys, 'frozen', False))
    except Exception as e:
        print(f"\n启动失败: {e}")
        input("按回车键退出...")